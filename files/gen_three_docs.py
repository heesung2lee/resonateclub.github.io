#!/usr/bin/env python3
"""Generate 3 Word documents for Resonate Club Inc. corporate setup."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

OUTDIR = "/Users/hslee/workspace/resonateclub.github.io/files/"

def set_cell_shading(cell, color):
    """Set cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    tcPr.append(shading)

def add_table(doc, headers, rows, col_widths=None):
    """Add a styled table to document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '1a3450')

    # Data rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r+1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if r % 2 == 0:
                set_cell_shading(cell, 'f0f4f8')

    doc.add_paragraph()
    return table

def set_narrow_margins(doc):
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x34, 0x50)
    return h

def add_body(doc, text):
    p = doc.add_paragraph(text)
    style = p.style
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(6)
    return p

# ══════════════════════════════════════════════════════════════════════
# DOCUMENT 1: US Investor Guide (English)
# ══════════════════════════════════════════════════════════════════════

def create_doc1_us_investor():
    doc = Document()
    set_narrow_margins(doc)

    # Title page
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("RESONATE CLUB INC.\n(리조네이트 클럽 주식회사)")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1a, 0x34, 0x50)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle.add_run("U.S. Investor Comprehensive Guide\n— Corporate Establishment, Visa & Investment Framework —")
    run2.font.size = Pt(13)
    run2.font.color.rgb = RGBColor(0xd4, 0xaf, 0x37)

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_text = f"Target Establishment Date: August 1, 2026\nInvestor: U.S. Citizen/Entity\nOwnership: 90% Equity Stake\nPrepared: {datetime.date.today().strftime('%B %d, %Y')}"
    info.add_run(info_text).font.size = Pt(10)

    doc.add_page_break()

    # ── EXECUTIVE SUMMARY ──
    add_heading_styled(doc, "Executive Summary", 1)
    add_body(doc,
        "This document outlines everything a U.S. investor needs to know and do to establish "
        "Resonate Club Inc. (리조네이트 클럽 주식회사), a premium Korea↔USA travel corporation, "
        "by the target date of August 1, 2026.\n\n"
        "The investor will own 90% equity (1,800 shares out of 2,000). Hee Sung Lee (F-4 visa holder) "
        "will serve as Representative Director (CEO) and own 10% equity (200 shares). The corporation "
        "will be registered as a Domestic & International Travel Agency (국내외여행업 / 기획여행사) "
        "in Haeundae-gu, Busan, under the Foreign Investment Promotion Act (FIPA).\n\n"
        "The investor's minimum cash requirement is approximately KRW 135 million (~USD $100,000) "
        "covering paid-in capital (KRW 100M) and initial setup costs (KRW 35M). The investor qualifies "
        "for a D-8-1 Corporate Investment Visa with full residency rights, family accompaniment, "
        "tax incentives, and a pathway to permanent residency (F-5)."
    )

    # ── 1. INVESTMENT STRUCTURE ──
    add_heading_styled(doc, "1. Investment Structure & Ownership", 1)

    add_table(doc,
        ["Item", "Detail"],
        [
            ["Legal Entity", "Resonate Club Inc. (리조네이트 클럽 주식회사) — Joint-Stock Company (주식회사)"],
            ["Brand", "Resonate Tour (결 투어)"],
            ["Registered Office", "Busan, Haeundae-gu, Republic of Korea"],
            ["Business Type", "Domestic & International Travel Agency (국내외여행업) — Planned Travel (기획여행사)"],
            ["Total Shares", "2,000 common shares (par value KRW 5,000/share)"],
            ["Issue Price", "KRW 50,000/share (total capitalization KRW 100,000,000)"],
            ["Investor (90%)", "1,800 shares / KRW 90,000,000 — Registered as Director or Auditor"],
            ["CEO / F-4 Holder (10%)", "200 shares / KRW 10,000,000 — Representative Director (대표이사)"],
            ["Total Paid-in Capital", "KRW 100,000,000 (~USD $74,000)"],
        ]
    )

    # ── 2. TIMELINE ──
    add_heading_styled(doc, "2. Master Timeline — Now Through August 1, 2026", 1)

    add_table(doc,
        ["Phase", "Key Dates", "Actions", "Who"],
        [
            ["Pre-Setup", "Now – Jul 10", "• Sign investor agreement / term sheet\n• Open Korean bank account (investor)\n• Prepare apostilled documents (US)", "Investor + CEO"],
            ["CEO Arrival", "Jul 17, 2026", "• CEO (Hee Sung Lee) arrives Korea on F-4 visa\n• Secure Haeundae-gu office lease\n• Prepare articles of incorporation", "CEO"],
            ["FDI Reporting", "Jul 20–25", "• File Foreign Investment Notification (외국인투자신고) at KOTRA or foreign exchange bank\n• Transfer investment funds to escrow/capital account\n• Obtain FDI registration certificate", "Investor + CEO"],
            ["Corp. Registration", "Jul 25–Aug 5", "• Court registration (부산지방법원 동부지원)\n• Corporate seal certificate issuance\n• Business registration (해운대세무서)\n• Open corporate bank account", "CEO (with notarized investor POA)"],
            ["Travel License", "Aug 5–Sep 5", "• Submit travel agency application (해운대구청)\n• Secure guarantee insurance (SGI 서울보증보험)\n• Complete registration — legally operational!", "CEO"],
            ["D-8 Visa", "Sep 5–Oct 5", "• Investor applies for D-8-1 visa (Korean Embassy in US)\n• Receives visa → enters Korea → ARC registration\n• Done! Investor has Korean residency", "Investor"],
        ]
    )

    # ── 3. D-8 VISA ──
    add_heading_styled(doc, "3. D-8-1 Corporate Investment Visa — Complete Brief", 1)

    add_heading_styled(doc, "3.1 Visa Overview", 2)
    add_table(doc,
        ["Item", "Detail"],
        [
            ["Visa Type", "D-8-1 (Corporate Investment)"],
            ["Eligibility", "Foreign national investing ≥ KRW 100M in Korean corporation, owning ≥ 10% voting shares"],
            ["Investment Requirement", "KRW 100,000,000 minimum — your KRW 90M investment qualifies if combined with total capital of KRW 100M+ and 10%+ share ownership"],
            ["Validity", "5 years per issuance, renewable indefinitely"],
            ["Work Scope", "Only for the invested company — cannot take outside employment"],
        ]
    )

    add_heading_styled(doc, "3.2 Required Documents (Investor)", 2)
    add_body(doc,
        "• Completed visa application form\n"
        "• Passport (valid 6+ months) + photo (3.5×4.5cm)\n"
        "• FDI Registration Certificate (외국인투자기업등록증명서)\n"
        "• Corporate registry extract (법인등기부등본)\n"
        "• Business registration certificate (사업자등록증)\n"
        "• Shareholder registry (주주명부)\n"
        "• Articles of incorporation (정관)\n"
        "• Bank certificate proving fund transfer (입금증명서)\n"
        "• Business plan (English accepted)\n"
        "• Office lease agreement (실제 사무실 필수)\n"
        "• Criminal background check — apostilled (FBI background check + US Dept. of State apostille)\n"
        "• Visa fee: ~KRW 60,000–120,000"
    )

    add_heading_styled(doc, "3.3 Application Process", 2)
    add_body(doc,
        "Overseas Route (Recommended):\n"
        "  1. Corporation established in Korea (CEO handles)\n"
        "  2. FDI registration completed (KOTRA/bank)\n"
        "  3. Investment funds transferred to corporate account\n"
        "  4. Investor applies at Korean Embassy/Consulate in U.S. (5–10 business days)\n"
        "  5. Visa issued → enter Korea → apply for Alien Registration Card (ARC) within 90 days\n\n"
        "Domestic Route (Alternative):\n"
        "  1. Enter Korea on visa-free/ESTA (90 days for U.S. citizens)\n"
        "  2. Complete all corporate setup within Korea\n"
        "  3. Apply for D-8 status change at local Immigration Office (2–4 weeks)\n"
        "  4. Receive ARC\n\n"
        "NOTE: FDI registration must occur BEFORE funds transfer. Doing it out of order triggers tax penalties."
    )

    add_heading_styled(doc, "3.4 Investor Perks & Benefits", 2)
    add_table(doc,
        ["Category", "Benefit", "Details"],
        [
            ["Residency", "5-Year Renewable Visa", "Long-term stable residence; no exit/re-entry permits needed"],
            ["Family", "F-3 Dependent Visa", "Spouse + minor children eligible for same-duration visas"],
            ["Health", "NHIS (National Health Insurance)", "Full Korean healthcare coverage — equivalent to citizens"],
            ["Tax", "Corporate Tax Reduction (FIPA)", "5–7 years reduced corporate income tax for foreign-invested enterprises"],
            ["Tax", "Local Tax Exemption", "Acquisition tax, property tax exemptions up to 15 years"],
            ["Tax", "Customs Duty Exemption", "Capital goods imported for the business exempt from duties"],
            ["Tax", "US-Korea Tax Treaty", "Avoidance of double taxation on dividends, interest, royalties"],
            ["Banking", "Corporate Banking + FX", "Full corporate bank accounts; foreign currency profit/dividend remittance permitted"],
            ["PR Path", "F-5 Permanent Residency", "After 5 years on D-8 + KRW 300M investment OR 2+ Korean employees + income ≥ GNI per capita + clean record"],
            ["Mobility", "Multiple Entry", "Unlimited entry/exit during visa validity"],
        ]
    )

    # ── 4. FINANCIAL OVERVIEW ──
    add_heading_styled(doc, "4. Financial Overview", 1)

    add_heading_styled(doc, "4.1 Capital Requirements", 2)
    add_table(doc,
        ["Category", "Amount (KRW)", "Amount (USD approx.)", "Notes"],
        [
            ["Paid-in Capital (100%)", "100,000,000", "$74,000", "Minimum for 국내외여행업 per Tourism Promotion Act"],
            ["Investor 90%", "90,000,000", "$66,700", "1,800 shares at KRW 50,000/share"],
            ["CEO 10%", "10,000,000", "$7,400", "200 shares — from personal funds"],
            ["Initial Setup Costs", "~5,000,000", "$3,700", "Registration fees, notary, office deposit, equipment"],
            ["1-Year Operating Reserve", "~20,000,000", "$14,800", "Monthly fixed costs × 12 + contingency"],
            ["TOTAL Recommended", "~125,000,000", "~$92,600", "Covers everything for first year"],
        ]
    )

    add_heading_styled(doc, "4.2 Monthly Fixed Costs (Operating)", 2)
    add_table(doc,
        ["Item", "Monthly (KRW)", "Annual (KRW)", "Note"],
        [
            ["Office Rent (Haeundae-gu)", "800,000", "9,600,000", "Small office, premium area"],
            ["Guarantee Insurance (SGI)", "83,000", "1,000,000", "0.5–1% of KRW 200M bond"],
            ["Accounting & Tax Services", "300,000", "3,600,000", "Monthly bookkeeping + annual tax filing"],
            ["Telecom & Office Supplies", "200,000", "2,400,000", "Internet, phone, sundries"],
            ["Miscellaneous", "200,000", "2,400,000", "Bank fees, memberships, etc."],
            ["Monthly Total", "1,583,000", "19,000,000", ""],
        ]
    )

    add_heading_styled(doc, "4.3 Revenue Model", 2)
    add_body(doc,
        "Core Model: Cost-Plus 40% Markup\n\n"
        "Primary Service Line — Korea→USA Premium Golf Tours:\n"
        "• Trump Premier (6N/7D, 4 rounds): Cost ~$6,966 → Sell $9,752/person\n"
        "• Trump Elite (5N/6D, 3 rounds): Cost ~$5,529 → Sell $7,741/person\n"
        "• Small private groups: 2–8 persons\n\n"
        "Year 1 Conservative Target (starting Sep 2026):\n"
        "• 12 teams × avg 4 persons × ~$2,500 profit/person ≈ $120,000 gross profit\n"
        "• After operating costs: ~KRW 130M (~$96,000) net operating profit\n\n"
        "Year 2 (full calendar):\n"
        "• 24 teams → ~KRW 312M (~$231,000) net profit\n\n"
        "Future: USA→Korea inbound luxury tours (requires KRW 300M capital upgrade to 종합여행업)"
    )

    # ── 5. KEY DECISIONS ──
    add_heading_styled(doc, "5. Key Decisions for Investor", 1)
    add_body(doc,
        "Before proceeding, the investor should confirm:\n\n"
        "1. INVESTMENT AMOUNT: Is the total investment KRW 125M (~$93K) or a higher amount?\n"
        "   (More capital = stronger financials for SGI guarantee insurance approval)\n\n"
        "2. ROLE: Will the investor be a passive director/auditor, or actively involved in operations?\n"
        "   (Passive role = standard; active role requires D-8 visa + physical presence)\n\n"
        "3. FAMILY: Does spouse/children need F-3 dependent visas?\n"
        "   (F-3 spouses cannot work in Korea — separate work visa would be needed)\n\n"
        "4. PR PATHWAY: Is permanent residency (F-5) a goal?\n"
        "   (KRW 300M investment OR hiring 2+ Korean employees opens the F-5-5 route after 5 years)\n\n"
        "5. TAX STRUCTURE: Consult a cross-border tax advisor for optimal US-Korea structure.\n"
        "   (Korean corporation tax + US foreign tax credit interaction needs planning)\n\n"
        "6. TIMING: Confirm the August 1 target is realistic.\n"
        "   (Corporation registration by Aug 1 is achievable; full licensing + D-8 visa extends to Oct)"
    )

    # ── 6. RISK ──
    add_heading_styled(doc, "6. Risk Considerations", 1)
    add_table(doc,
        ["Risk", "Level", "Mitigation"],
        [
            ["Tourism Promotion Act amendment (capital increase)", "Medium", "Investor-backed capital already exceeds likely requirements"],
            ["SGI guarantee insurance denial", "Medium", "Strong business plan + 신용보증기금 route as backup"],
            ["D-8 visa processing delays", "Low", "Start early; embassy processing typically 5–10 business days"],
            ["Currency fluctuation (KRW/USD)", "Low", "Revenue in USD, costs split KRW/USD — natural hedge"],
            ["Market adoption / slow first season", "Medium", "Conservative Year 1 projections; low fixed cost base"],
            ["Korea political/regulatory changes", "Low", "Travel industry is well-established; FIPA protects foreign investors"],
            ["US tax compliance penalties", "High", "Form 5471 non-filing = $10,000+ penalty per form per year — engage CPA before investing"],
            ["PFIC classification risk", "Low", "Active travel business should NOT be PFIC, but confirm with CPA; Form 8621 filing is punitive"],
        ]
    )

    # ── 7. CROSS-BORDER TAX PLANNING ═══════════════════════════════
    add_heading_styled(doc, "7. Cross-Border Tax Planning & Compliance Timeline", 1)
    add_body(doc,
        "This section outlines EVERY tax-related action the investor needs to take — in both the "
        "United States and Korea — organized chronologically from pre-investment through ongoing "
        "annual compliance. Engaging a qualified cross-border CPA (U.S. side) and a Korean 세무사 "
        "(tax accountant) is NOT optional — penalties for non-compliance are severe on both sides."
    )

    # 7.1
    add_heading_styled(doc, "7.1 Tax Professionals to Engage", 2)
    add_table(doc,
        ["Professional", "Jurisdiction", "When to Engage", "Estimated Cost", "Key Role"],
        [
            ["U.S. Cross-Border CPA\n(EA or CPA with intl experience)", "USA (IRS)", "NOW — before any money moves", "$2,500–5,000 initial\n$1,500–3,000/year", "Form 5471, FBAR, FATCA, foreign tax credit (Form 1116), structure advice, PFIC analysis"],
            ["Korean 세무사\n(Tax Accountant)", "Korea (NTS)", "Jul–Aug 2026 (before 사업자등록)", "₩300,000–500,000/month\n(ongoing bookkeeping)", "Corporate tax filing, VAT returns, payroll tax (if hiring), FIPA incentive applications, withholding tax on dividends"],
            ["U.S. Estate Planning Attorney\n(if significant net worth)", "USA", "Before investing", "$3,000–7,000 one-time", "Korean situs assets in U.S. estate, gift tax implications of share transfers, QDOT planning if non-citizen spouse"],
            ["Korean 법무사\n(Legal Administrator)", "Korea", "Jul 2026", "₩500,000–800,000", "Corporation registration, FDI filing, shareholder registry maintenance"],
        ]
    )

    # 7.2
    add_heading_styled(doc, "7.2 Pre-Investment Tax Structuring — NOW Through July 15, 2026", 2)
    add_body(doc,
        "⚠️ CRITICAL: The decisions made BEFORE money moves are the most consequential for long-term tax outcomes.\n\n"
        "Action 1 — Ownership Structure Decision (NOW):\n"
        "• Option A: Direct individual ownership (investor name directly on Korean shareholder registry)\n"
        "  → Simpler setup; Form 5471 Category 3 filer; dividends taxed at qualified dividend rates if eligible\n"
        "• Option B: U.S. LLC/S-Corp as holding entity\n"
        "  → Added complexity & cost; potential state tax benefits; Form 5471 Category 4 or 5 filer\n"
        "  → S-Corp CANNOT have foreign shareholders; if investor has non-US family, use LLC or C-Corp\n"
        "• Recommendation: Start with direct ownership (Option A) unless estate planning or anonymity concerns dictate otherwise.\n\n"
        "Action 2 — PFIC Analysis (NOW):\n"
        "• A Passive Foreign Investment Company (PFIC) triggers Form 8621 — one of the most punitive IRS regimes\n"
        "• Test: >50% passive income OR >50% passive assets = PFIC\n"
        "• Resonate Club is an ACTIVE travel business — should NOT be PFIC\n"
        "• GET CPA OPINION LETTER confirming non-PFIC status — do NOT skip this\n\n"
        "Action 3 — Investment Tracking System Setup (NOW):\n"
        "• Set up spreadsheet/software to track: investment date, KRW amount, USD exchange rate at each transfer\n"
        "• Needed for: foreign tax credit calculation, future exit capital gains basis, Section 986(c) exchange rate elections\n\n"
        "Action 4 — FBAR/Form 8938 Readiness (NOW):\n"
        "• FBAR (FinCEN Form 114): Required if aggregate foreign accounts >$10,000 at any point in calendar year\n"
        "• Form 8938 (FATCA): Required if foreign financial assets >$200,000 (single) / $400,000 (married) on Dec 31\n"
        "• Both are INFORMATION returns (not tax due) — but penalties are $10,000+ per violation\n"
        "• Plan for the Korean corporate account(s) + any personal Korean accounts to trigger these thresholds\n\n"
        "Action 5 — Capitalization Strategy (consult CPA):\n"
        "• Equity vs. shareholder loan mix: loans can be repaid tax-free (vs. dividends subject to withholding)\n"
        "• Thin capitalization rules in Korea: interest deductions may be limited if debt-to-equity ratio > 2:1\n"
        "• Recommendation: At least 70% equity, 30% shareholder loan if debt structure desired"
    )

    # 7.3
    add_heading_styled(doc, "7.3 Korean FIPA Tax Incentives — Timeline & Application Process", 2)

    add_body(doc,
        "Under the Foreign Investment Promotion Act (FIPA / 외국인투자촉진법), foreign-invested enterprises "
        "qualify for the following tax benefits. Each requires proactive APPLICATION — they are not automatic."
    )

    add_table(doc,
        ["Incentive", "Benefit", "Duration", "When to Apply", "Who Handles It"],
        [
            ["Corporate Income Tax\nReduction (법인세 감면)", "100% exemption years 1-5\n50% reduction years 6-7\n(Applies only to income from the qualified business)", "7 years", "Within FY of first profit\n(or by tax filing deadline)", "Korean 세무사 — files with NTS\n(FIPA Article 121-2, RETI Article 116-2)"],
            ["Acquisition Tax\nExemption (취득세 면제)", "100% exemption on real estate,\nvehicles, equipment acquired\nfor the business", "15 years from\nbusiness registration", "At time of each\nqualifying purchase", "Korean 세무사 — files with local\ngovernment (지방자치단체)\n(FIPA Article 121-3)"],
            ["Property Tax\nReduction (재산세 감면)", "100% years 1-5\n50% years 6-15\n(on business property)", "15 years", "Annual — with property\ntax filing (July)", "Korean 세무사"],
            ["Customs Duty Exemption\n(관세 면제)", "100% exemption on capital\ngoods imported for the business\n(e.g., tour vehicles, camera\nequipment, golf simulators)", "Within 3 years of\nFDI registration", "At each import — file\nwith Korea Customs\nService (관세청)", "Korean 세무사/관세사\nPre-approval: HS Code\n사전심사 recommended"],
            ["Dividend Withholding\nTax Reduction", "Under US-Korea Tax Treaty:\n15% standard → 10% if 10%+\nownership → possibly 5%", "Ongoing", "When declaring dividends", "Korean 세무사 — reduced\nrate application + U.S. Form\n8833 (treaty-based return\nposition disclosure)"],
        ]
    )

    add_body(doc,
        "⚠️ CRITICAL for FIPA Benefits:\n"
        "• Must file 외국인투자기업 등록신청 (FDI Enterprise Registration) BEFORE any tax benefit application\n"
        "• Tax reduction application (조세감면신청서) must be filed within the first taxable year after earning profit\n"
        "• MISSING the application deadline = permanent loss of the benefit for that year (no retroactive remedy)\n"
        "• The 세무사 must reference FIPA Article 121-2 through 121-5 and Restriction of Special Taxation Act (조세특례제한법)\n"
        "• Local tax exemptions require separate application to Busan Metropolitan City / Haeundae-gu office"
    )

    # 7.4
    add_heading_styled(doc, "7.4 US-Korea Tax Treaty — Key Provisions", 2)
    add_table(doc,
        ["Income Type", "Without Treaty", "With Treaty", "IRS Form"],
        [
            ["Dividends (배당소득)", "30% withholding", "15% (standard)\n10% if ≥10% ownership\n5% if ≥10% ownership\n+ certain conditions", "IRS Form 1116\n(Foreign Tax Credit)\nForm 8833 (Treaty Disclosure)"],
            ["Interest (이자소득)", "30% withholding", "12% (standard)\nMay be reduced/exempt\nfor certain loans", "Form 1116"],
            ["Royalties (사용료)", "30% withholding", "15% (standard)\n10% for certain IP", "Form 1116"],
            ["Capital Gains (양도소득)", "May be taxed\nin Korea", "Generally only taxed in\ncountry of residence\n(unless real property)", "Form 1116 + Schedule D\n(if US-situs)"],
            ["Branch Profits", "20%+ branch tax", "N/A — Korean corporation\nis separate entity", "N/A"],
        ]
    )

    add_body(doc,
        "Key Points:\n"
        "• Korean withholding tax on dividends (10-15%) is CREDITABLE against U.S. tax liability via Form 1116\n"
        "• If U.S. tax > Korean tax: pay the difference to IRS. If Korean tax > U.S. tax: carry forward excess credits\n"
        "• Form 8833 must be filed with U.S. return when claiming treaty benefits (generally required for ≥10% ownership)\n"
        "• Treaty benefit limitations: must be a \"qualified resident\" of the U.S. under Article 16 (Limitation on Benefits)\n"
        "• U.S. citizens are ALWAYS taxed on worldwide income regardless of residence — treaty does not override this"
    )

    # 7.5
    add_heading_styled(doc, "7.5 Annual U.S. Compliance Obligations (Ongoing)", 2)
    add_table(doc,
        ["Form", "Name", "Threshold", "Deadline", "Penalty for Non-Compliance"],
        [
            ["Form 5471", "Information Return of U.S. Persons\nwith Respect to Certain Foreign\nCorporations", "10%+ ownership of a foreign\ncorporation (Category 3 filer\nfor 90% shareholder)", "Due with tax return\n(Apr 15 + extensions)", "$10,000 per form per year;\nadditional $10,000–50,000\nfor continued failure\nafter IRS notice"],
            ["FBAR\n(FinCEN 114)", "Report of Foreign Bank\nand Financial Accounts", "Aggregate foreign accounts\n>$10,000 at any time", "April 15 (auto-extension\nto Oct 15)", "$10,000+ per violation\n(if non-willful); greater of\n$100,000 or 50% of account\nbalance if willful"],
            ["Form 8938", "Statement of Specified\nForeign Financial Assets\n(FATCA)", ">$200,000 on Dec 31\n(or >$300,000 at any time)\nSingle filer thresholds", "Due with tax return", "$10,000 initial + $10,000\nper 30 days (max $60,000)"],
            ["Form 1116", "Foreign Tax Credit", "Any foreign taxes paid\nor accrued", "Due with tax return", "Loss of credit (pay double tax\non same income)"],
            ["Form 8833", "Treaty-Based Return\nPosition Disclosure", "Generally required when\nclaiming treaty benefits with\n≥10% ownership", "Due with tax return", "Potential $1,000 penalty\n(rarely enforced, but\nfailure = weaker audit\nposition)"],
            ["Form 926", "Return by a Transferor of\nProperty to a Foreign\nCorporation", "Transfer of property\n(including cash) to foreign\ncorporation", "Due with tax return", "$10,000+ per failure"],
        ]
    )

    add_body(doc,
        "⚠️ IMPORTANT NOTES:\n"
        "• Form 5471 is the HIGHEST RISK compliance item for a 90% shareholder\n"
        "  - Section 6038 requires detailed financial reporting of the Korean corporation\n"
        "  - Korean GAAP financials must be translated/reconciled to U.S. GAAP\n"
        "  - Schedule G (income statement), Schedule H (balance sheet), Schedule M (transactions)\n"
        "  - The penalty of $10,000 applies EVEN IF NO TAX IS DUE\n\n"
        "• The Korean corporation must provide the following annually to the investor's CPA:\n"
        "  - Full financial statements (K-GAAP or IFRS) + Korean tax return\n"
        "  - Shareholder registry confirming ownership percentage\n"
        "  - All related-party transaction details (investor ↔ corporation)\n"
        "  - Details of any capital increases, share transfers, or structural changes\n\n"
        "• Exchange Rate:\n"
        "  - Use year-end rate (Dec 31) for balance sheet items\n"
        "  - Use average annual rate for income statement items\n"
        "  - Track the exchange rate at each capital contribution date for basis purposes"
    )

    # 7.6
    add_heading_styled(doc, "7.6 Integrated Tax Compliance Timeline", 2)
    add_table(doc,
        ["When", "Action", "Who", "Deliverable / Result"],
        [
            ["NOW\n(Jun 2026)", "① Engage U.S. cross-border CPA\n② PFIC analysis opinion letter\n③ Ownership structure decision\n④ Open tracking spreadsheet for basis/FX", "Investor + CPA", "CPA engagement letter;\nPFIC non-classification opinion;\nFX tracking sheet created"],
            ["Jul 15, 2026", "⑤ Sign shareholder agreement\n⑥ Prepare source-of-funds documentation", "Investor + CEO", "Signed agreement;\nbank statements, tax returns"],
            ["Jul 20–25, 2026\n(FDI Stage)", "⑦ File foreign investment notification (KOTRA)\n⑧ Transfer investment funds (track USD-KRW rate!)\n⑨ Obtain FDI registration certificate", "CEO (with investor POA)", "FDI certificate;\nBank remittance receipt\nwith exchange rate recorded"],
            ["Aug 1–6, 2026\n(Corp. Reg.)", "⑩ Register corporation\n⑪ Open corporate bank account\n⑫ Engage Korean 세무사", "CEO + 법무사", "법인등기부등본;\n세무사 engagement letter"],
            ["Aug 7–11, 2026\n(Biz Reg.)", "⑬ Register for business (사업자등록)\n⑭ Register for VAT\n⑮ Apply for corporate tax ID (법인세 납세번호)", "CEO + 세무사", "사업자등록증;\nNTS taxpayer number"],
            ["Sep 2026", "⑯ Apply for FIPA corporate tax reduction\n⑰ Apply for local tax exemptions (취득세/재산세)\n⑱ File FIPA incentive application (조세감면신청서)", "세무사", "FIPA tax benefit\napproval letter from NTS;\nLocal government exemption certificates"],
            ["Oct 2026 onward\n(upon first asset purchase)", "⑲ For each qualifying capital goods import:\n  - Pre-file HS Code 사전심사 with 관세청\n  - Apply for customs duty exemption\n  - Keep all import documentation", "세무사 + 관세사", "Customs duty exemption;\nAsset register maintained"],
            ["Mar 2027\n(First Korean tax return)", "⑳ File first corporate tax return (법인세 신고)\n㉑ Claim FIPA tax reduction on first profit\n㉒ File VAT returns (부가세 신고 — quarterly)", "세무사", "Corporate tax return;\nVAT returns"],
            ["Apr 15, 2027\n(First US tax return with\nforeign holdings)", "㉓ Prepare Korean GAAP→US GAAP reconciliation\n㉔ File Form 5471 (Category 3) with U.S. return\n㉕ File FBAR (FinCEN 114)\n㉖ File Form 1116 (Foreign Tax Credit)\n㉗ File Form 8938 (FATCA)\n㉘ File Form 926 (if applicable)\n㉙ File Form 8833 (Treaty Disclosure)", "Investor + U.S. CPA", "U.S. tax return with all\ninternational information\nforms attached;\nFBAR acknowledgment"],
            ["Annually thereafter\n(Apr 15 IRS / Mar 31 NTS)", "㉚ Repeat all U.S. reporting (5471, FBAR, 8938, 1116)\n㉛ Korean corporate tax return + VAT\n㉜ Monitor FIPA benefit expiration (Year 6: corporate tax\n   drops to 50%; Year 8: ends entirely)", "CPA + 세무사\n+ Investor", "Annual tax filings;\nFIPA benefit tracking sheet"],
            ["Upon Dividend\nDeclaration", "㉝ Korean withholding tax at treaty rate (10-15%)\n㉞ Korean 세무사 files reduced rate application\n㉟ U.S. CPA claims foreign tax credit (Form 1116)", "세무사 + CPA", "Dividend payment with\ntreaty-reduced withholding;\nForm 1116 credit claimed"],
            ["Upon Future\nExit / Sale", "㊱ Korean capital gains treatment (if any)\n㊲ U.S. capital gains tax on sale of foreign corp shares\n㊳ Basis calculation using tracked FX rates\n㊴ Exit tax planning with CPA 12+ months before sale", "CPA + 세무사\n+ Investor", "Exit tax strategy;\nCapital gains calculation"],
        ]
    )

    # 7.7
    add_heading_styled(doc, "7.7 Dividend & Profit Repatriation Strategy", 2)
    add_body(doc,
        "How the investor eventually takes money out of Korea matters significantly for tax.\n\n"
        "Route A — Dividend Distribution (배당):\n"
        "• Korean corp declares dividend → 10–15% Korean withholding tax → investor receives net\n"
        "• U.S. tax: dividend included in gross income; foreign tax credit for Korean withholding via Form 1116\n"
        "• Net U.S. rate: 20% qualified dividend rate – 15% credit = 5% net additional U.S. tax (approx.)\n"
        "• Best when: profits are significant and reinvestment is not needed\n\n"
        "Route B — Shareholder Loan Repayment (if structured upfront):\n"
        "• Principal repayment: tax-free (no withholding)\n"
        "• Interest: subject to 12% Korean withholding (treaty rate), creditable in U.S.\n"
        "• Best when: investor provided some capital as debt (see 7.2 § Action 5)\n\n"
        "Route C — Management/Consulting Fee:\n"
        "• Korean corp pays fee to U.S. investor for legitimate services (director oversight, strategic consulting)\n"
        "• Korean deduction for corp; U.S. ordinary income for investor\n"
        "• Must be at arm's length (transfer pricing documentation required)\n"
        "• Best when: investor provides genuine, documentable services\n\n"
        "Route D — Salary (if investor holds D-8 visa and physically works in Korea):\n"
        "• Korean salary tax + U.S. income tax (foreign earned income exclusion possible via Form 2555)\n"
        "• Best when: investor physically resides in Korea and actively manages the business\n\n"
        "⚠️ WARNING: ALL repatriation routes require documentation. Korean tax authority (NTS) "
        "aggressively audits related-party transactions involving foreign shareholders. Transfer pricing "
        "rules apply. Maintain board resolutions, invoices, time logs, and arm's-length pricing analyses."
    )

    # ── 8. CONTACTS ──
    add_heading_styled(doc, "8. Key Contacts & Resources", 1)
    add_table(doc,
        ["Contact", "Role", "Website / Phone"],
        [
            ["KOTRA Invest KOREA", "FDI registration + incentives", "investkorea.org / +82-1600-7119"],
            ["Busan District Court (East)", "Corporation registration", "busan.scourt.go.kr"],
            ["Haeundae Tax Office", "Business registration + corporate tax", "nts.go.kr / +82-51-740-2000"],
            ["Haeundae-gu Office", "Travel agency license", "haeundae.go.kr / +82-51-749-4000"],
            ["Seoul Guarantee Insurance (SGI)", "Guarantee insurance bond", "sgi.co.kr / +82-1588-7000"],
            ["Korea Immigration Service", "D-8 visa / ARC", "hikorea.go.kr / 1345"],
            ["Korean Embassy in US", "Visa application", "overseas.mofa.go.kr"],
            ["Korea Customs Service (관세청)", "HS Code pre-review + customs exemption", "customs.go.kr / +82-1577-8577"],
            ["National Tax Service (국세청)", "Corporate tax filing, FIPA incentives", "nts.go.kr / +82-126"],
            ["IRS — International Tax", "Form 5471, FBAR, FATCA questions", "irs.gov / +1-267-941-1000 (Intl)"],
            ["Resonate Club Inc. (CEO)", "Hee Sung Lee", "resonatetour.com"],
        ]
    )

    add_body(doc, "\n---\nThis document is a comprehensive guide based on publicly available information as of June 2026. "
              "The investor should consult a Korean-licensed legal professional (법무사/변호사), a U.S. cross-border "
              "CPA, and a Korean 세무사 before executing the investment. Tax laws, treaty provisions, and FIPA benefit "
              "rules are subject to change. Penalties for U.S. international information return non-compliance are severe "
              "($10,000+ per form per year) regardless of whether any tax is ultimately due.")

    path1 = OUTDIR + "Resonate_Club_US_Investor_Guide_EN.docx"
    doc.save(path1)
    print(f"[OK] Document 1 saved: {path1}")
    return path1


# ══════════════════════════════════════════════════════════════════════
# DOCUMENT 2: F-4 CEO Action Plan (Korean)
# ══════════════════════════════════════════════════════════════════════

def create_doc2_ceo_actionplan():
    doc = Document()
    set_narrow_margins(doc)

    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("리조네이트 클럽 주식회사\nF-4 대표이사 실행계획")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1a, 0x34, 0x50)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle.add_run("2026년 7월 17일 입국 기준 — 단계별 액션플랜")
    run2.font.size = Pt(12)
    run2.font.color.rgb = RGBColor(0xd4, 0xaf, 0x37)

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f"작성일: 2026년 6월 11일\n대표이사: 이희성 (F-4 재외동포 비자)\n투자자: 미국인 (D-8-1 예정, 지분 90%)").font.size = Pt(10)

    doc.add_page_break()

    # ── 개요 ──
    add_heading_styled(doc, "개요", 1)
    add_body(doc,
        "이 문서는 2026년 7월 17일 한국에 입국하는 F-4 비자 소지자 이희성 대표이사가 "
        "8월 1일까지 법인설립을 완료하고, 이후 여행업 등록 및 모든 인허가를 완료하기까지의 "
        "모든 단계를 시간순으로 정리한 실행 매뉴얼입니다.\n\n"
        "핵심 전제:\n"
        "• 투자자: 미국인, 자본금 9,000만원 출자 (지분 90%), D-8-1 비자 취득 예정\n"
        "• 대표이사: 이희성, F-4 비자, 자본금 1,000만원 출자 (지분 10%)\n"
        "• 법인형태: 외국인투자기업 주식회사 (FIPA 적용)\n"
        "• 사업유형: 국내외여행업 (기획여행사)\n"
        "• 소재지: 부산광역시 해운대구\n"
        "• 목표: 8월 중순까지 여행업 등록 완료 → 9월부터 영업 개시"
    )

    # ── TIME ESTIMATE SUMMARY ──
    add_heading_styled(doc, "전체 타임라인 요약", 1)
    add_table(doc,
        ["단계", "예상 소요 기간", "누적", "목표 완료일"],
        [
            ["사전 준비 (미국)", "입국 전 완료", "D-14", "7월 3일까지"],
            ["STEP 1: 사무실 계약", "1~3일", "D+3", "7월 20일"],
            ["STEP 2: 정관 작성·공증", "2~3일", "D+6", "7월 23일"],
            ["STEP 3: 외국인투자신고 + 자본금 입금", "3~7일", "D+13", "7월 30일"],
            ["STEP 4: 법인설립등기", "5~7일", "D+20", "8월 6일"],
            ["STEP 5: 사업자등록", "3~5일", "D+25", "8월 11일"],
            ["STEP 6: 국내외여행업 등록", "14~30일", "D+55", "9월 10일"],
            ["STEP 7: 영업보증보험 가입", "5~7일", "D+62", "9월 17일"],
            ["STEP 8: D-8 비자 투자자 지원", "30~60일", "D+120", "10월~11월"],
            ["영업 개시", "—", "—", "9월 (등록 완료 즉시)"],
        ]
    )

    # ── STEP-BY-STEP ──
    steps = [
        ("STEP 0: 입국 전 사전 준비 (현재 ~ 7월 16일)",
         "소요 기간: 입국 전까지\n\n"
         "1. 투자자와 주주간 계약서 체결\n"
         "   - 지분 비율: 투자자 90% / CEO 10%\n"
         "   - 자본금: 총 1억원 (투자자 9,000만원, CEO 1,000만원)\n"
         "   - 대표이사: 이희성 / 이사(또는 감사): 투자자\n"
         "   - 배당 정책, 의사결정 권한, 지분양도 제한 등 명시\n\n"
         "2. 투자자 서류 준비 (미국)\n"
         "   - 여권 사본 (6개월 이상 유효기간)\n"
         "   - FBI 범죄경력증명서 → 국무부 Apostille (D-8 비자용, 4~6주 소요)\n"
         "   - 투자자금 원천 증명 (은행 statement, 세금신고서 등)\n\n"
         "3. 한국 법무사/행정사 사전 섭외\n"
         "   - 부산 지역 법인등기 전문 법무사\n"
         "   - 예상 비용: 50~80만원 (법인설립 패키지)\n\n"
         "4. 사무실 임차 사전 리서치\n"
         "   - 해운대구 내 소형 오피스 (월 50~80만원)\n"
         "   - 부동산 중개업소 컨택, 매물 리스트 확보\n\n"
         "5. 외국인투자신고 사전 상담\n"
         "   - KOTRA 외국인투자종합상담센터(1600-7119) 전화 상담\n"
         "   - 외환은행 기업금융 지점 사전 방문 예약"),
        ("STEP 1: 사무실 계약 (7월 17~20일, 1~3일)",
         "소요 기간: 1~3일\n\n"
         "1. 입국 (7월 17일): 인천공항 → 부산 이동 (KTX 약 2시간40분)\n\n"
         "2. 해운대구 부동산 방문 (7월 18일)\n"
         "   - 사전 확보한 매물 리스트 중심으로 2~3곳 실사\n"
         "   - 필수 요건:\n"
         "     · 실제 사무공간 (가상오피스 불가 — 여행업/비자 심사에서 거절)\n"
         "     · 해운대구 관할 내 위치\n"
         "     · 확정일자 받을 수 있는 정식 임대차계약\n"
         "   - 예상 보증금: 500~1,000만원\n"
         "   - 예상 월세: 50~80만원\n\n"
         "3. 임대차 계약 체결 (7월 19~20일)\n"
         "   - 계약서 작성 + 확정일자 수령\n"
         "   - 전입신고는 대표이사 주소지로 (사무실 ≠ 주거)\n\n"
         "☑️ 확보해야 할 서류: 임대차계약서 사본, 확정일자, 사업장 도면"),
        ("STEP 2: 정관 작성 및 공증 (7월 21~23일, 2~3일)",
         "소요 기간: 2~3일\n\n"
         "1. 정관 작성\n"
         "   - 법무사 사무실 방문 또는 startbiz.go.kr 온라인 작성\n"
         "   - 필수 기재 사항:\n"
         "     · 상호: 리조네이트 클럽 주식회사 (영문: Resonate Club Inc.)\n"
         "     · 목적: 국내외여행업, 여행 알선, 관광 컨설팅 등\n"
         "     · 본점: 부산광역시 해운대구 [주소]\n"
         "     · 자본금: 1억원 (보통주 2,000주, 1주당 5만원)\n"
         "     · 발기인: 이희성 + [투자자명] (공동발기 또는 투자자 위임)\n"
         "     · 임원: 대표이사 1인, 이사(또는 감사) 1인\n\n"
         "2. 공증\n"
         "   - 부산지방법원 동부지원 인근 공증사무소\n"
         "   - 정관 + 발기인 인감증명서 준비\n"
         "   - 비용: 약 10~15만원\n\n"
         "☑️ 확보: 공증된 정관, 발기인 인감증명서"),
        ("STEP 3: 외국인투자신고 + 자본금 입금 (7월 24~30일, 3~7일)",
         "소요 기간: 3~7일\n\n"
         "⚠️ 중요: 반드시 외국인투자신고를 먼저 한 후, 자본금을 입금해야 합니다.\n"
         "순서가 바뀌면 과태료 및 세제혜택 상실 위험이 있습니다.\n\n"
         "1. 외국인투자신고 (7월 24~25일)\n"
         "   - 신고처: KOTRA 외국인투자종합상담센터 또는 외환은행 본점\n"
         "   - 제출 서류:\n"
         "     · 외국인투자신고서 (KOTRA 양식)\n"
         "     · 투자자 여권 사본\n"
         "     · 투자자 주소 증명서류\n"
         "     · 사업계획서\n"
         "     · 정관 (공증본)\n"
         "   - 처리 기간: 통상 1~2일 (현장 접수 시)\n\n"
         "2. 자본금 계좌 개설 (7월 25~26일)\n"
         "   - 외환은행 또는 국민은행 '외국인투자 전용 계좌' 개설\n"
         "   - 법인설립 전이므로 '발기인 명의의 자본금 납입 보관계좌'\n\n"
         "3. 자본금 입금 (7월 27~28일)\n"
         "   - 투자자: $67,000 → 원화 환전 후 9,000만원 입금\n"
         "   - CEO: 1,000만원 입금\n"
         "   - 은행에서 '주식납입금 보관증명서' 발급\n\n"
         "4. 외국인투자기업 등록 (7월 29~30일)\n"
         "   - 자본금 입금 확인 후 KOTRA/은행에서 등록증 발급\n"
         "   - 외국인투자기업등록증명서 수령 (D-8 비자 신청 필수 서류)\n\n"
         "☑️ 확보: 외국인투자기업등록증명서, 주식납입금 보관증명서"),
        ("STEP 4: 법인설립등기 (7월 31일 ~ 8월 6일, 5~7일)",
         "소요 기간: 관할 등기소 접수 후 5~7일\n\n"
         "1. 등기 신청서류 준비 (7월 31일 ~ 8월 1일)\n"
         "   - 법무사에게 서류 일임 추천 (패키지 50~80만원)\n"
         "   - 필요 서류:\n"
         "     · 정관 (공증본)\n"
         "     · 발기인 인감증명서 + 주민등록등본\n"
         "     · 주식납입금 보관증명서\n"
         "     · 대표이사 취임승낙서\n"
         "     · 이사(감사) 취임승낙서 (투자자)\n"
         "     · 등기신청서 + 등록면허세\n\n"
         "2. 등기소 접수 (8월 1~3일)\n"
         "   - 부산지방법원 동부지원 등기과 방문\n"
         "   - 또는 전자등기시스템(iros.go.kr) 이용\n"
         "   - 접수 후 처리 기간: 5~7일\n\n"
         "3. 등기완료 확인 (8월 6일 전후)\n"
         "   - 법인등기부등본 발급 (인터넷등기소에서 즉시 가능)\n"
         "   - 법인 인감카드 발급\n"
         "   - 법인 통장 개설 (은행 방문, 법인등기부등본 + 법인인감 필요)\n\n"
         "☑️ 확보: 법인등기부등본, 법인인감카드, 법인 통장"),
        ("STEP 5: 사업자등록 (8월 7~11일, 3~5일)",
         "소요 기간: 세무서 접수 후 3~5일\n\n"
         "1. 사업자등록 신청 (8월 7~8일)\n"
         "   - 해운대세무서 방문 또는 홈택스(hometax.go.kr) 온라인 신청\n"
         "   - 필요 서류:\n"
         "     · 법인등기부등본\n"
         "     · 임대차계약서 사본\n"
         "     · 정관 사본\n"
         "     · 법인인감증명서\n"
         "     · 대표이사 신분증\n\n"
         "2. 사업자등록증 수령 (8월 11일 전후)\n"
         "   - 업종: 여행사업 (관광진흥법) / 관광여행사\n"
         "   - 과세유형: 일반과세자\n"
         "   - 통상 3~5일 이내 발급 (온라인 시 더 빠름)\n\n"
         "☑️ 확보: 사업자등록증"),
        ("STEP 6: 국내외여행업 등록 (8월 12일 ~ 9월 10일, 14~30일)",
         "소요 기간: 구청 접수 후 2~4주 (구청 및 한국관광공사 심사)\n\n"
         "⚠️ 중요: 이 단계가 가장 긴 리드타임입니다. 서류 누락 시 반려되면 다시 2~4주입니다.\n\n"
         "1. 구청 제출 서류 준비 (8월 12~14일)\n"
         "   - 관광진흥법 시행규칙 제9조 서식:\n"
         "     · 여행업 등록신청서\n"
         "     · 사업계획서 (첨부 양식) — 아래 내용 포함 필수:\n"
         "       → 법인 개요, 사업 비전, 마케팅 전략, 재무 계획\n"
         "       → 취급 상품: 한국→미국 프리미엄 골프 투어\n"
         "       → 시장분석, 경쟁력 분석, 리스크 관리 방안\n"
         "     · 법인등기부등본\n"
         "     · 사업자등록증 사본\n"
         "     · 임대차계약서 사본\n"
         "     · 자본금 증명 서류\n"
         "     · 영업보증보험 가입 증명서 (동시 진행 가능)\n"
         "     · 국외여행 인솔자 자격증 사본 (또는 채용계획서)\n"
         "     · 수수료: 5만원 (수입증지)\n\n"
         "2. 해운대구청 접수 (8월 14~18일)\n"
         "   - 관광과/문화관광과 방문\n\n"
         "3. 심사 기간 (2~4주)\n"
         "   - 구청 1차 심사 → 필요 시 한국관광공사 기술 검토\n"
         "   - 추가 서류 요청 가능성 대비\n\n"
         "4. 등록증 수령 (9월 10일 전후)\n\n"
         "☑️ 확보: 국내외여행업 등록증"),
        ("STEP 7: 영업보증보험 가입 (8월 12일 ~ 9월 17일, 병행 진행)",
         "소요 기간: SGI 심사 5~7일 (STEP 6과 병행 가능)\n\n"
         "국내외여행업 + 기획여행(패키지) 기준 보증금액: 2억원\n\n"
         "1. 서울보증보험(SGI) 상담 예약 (8월 12일)\n"
         "   - 부산지점: 부산광역시 연제구 법원로 소재\n"
         "   - 필요 서류:\n"
         "     · 법인등기부등본\n"
         "     · 사업자등록증\n"
         "     · 사업계획서\n"
         "     · 재무제표 / 자본금 증명\n"
         "     · 대표이사 신용정보 동의서\n\n"
         "2. 보증보험 청약 (8월 14~15일)\n"
         "   - 방식 A: 전액 지급보증 (예치금 0원, 연 보험료 0.5~1.5%)\n"
         "     → 2억원 × 0.5~1% = 연 100~300만원\n"
         "   - 방식 B: 부분예치 + 지급보증 (심사 통과 어려울 시)\n"
         "     → 예: 4,000만원 예치 + 1.6억원 지급보증\n\n"
         "3. 보증서 발급 (8월 19~22일)\n\n"
         "☑️ 확보: 영업보증보험 증권 (구청 등록 제출용)"),
        ("STEP 8: 투자자 D-8 비자 지원 및 후속 조치 (9월 ~ )",
         "소요 기간: 30~60일\n\n"
         "9월부터 법인은 정상 영업 가능. 투자자 D-8 비자는 별도 진행.\n\n"
         "1. 투자자에게 필요 서류 송부\n"
         "   - 외국인투자기업등록증명서\n"
         "   - 법인등기부등본 (영문 번역본)\n"
         "   - 사업자등록증 (영문 번역)\n"
         "   - 주주명부\n"
         "   - 정관\n"
         "   - 사업계획서 (영문)\n"
         "   - 사무실 임대차계약서\n\n"
         "2. 투자자 → 주한미국대사관에서 D-8-1 비자 신청\n"
         "   - 처리 기간: 5~10 영업일\n\n"
         "3. 투자자 입국 → 외국인등록증(ARC) 발급\n"
         "   - 입국 후 90일 이내 관할 출입국사무소 방문\n\n"
         "4. 추가 등록 사항:\n"
         "   - 관광협회 가입 (한국관광협회중앙회)\n"
         "   - 국외여행 인솔자 등록 (한국관광공사)\n"
         "   - 배상책임보험 가입 (영업 전 필수)\n"
         "   - 여행자보험 상품 계약"),
    ]

    for i, (title, body) in enumerate(steps):
        add_heading_styled(doc, title, 1)
        for para_text in body.split("\n\n"):
            if para_text.strip():
                if para_text.strip().startswith("1.") or para_text.strip().startswith("2.") or \
                   para_text.strip().startswith("3.") or para_text.strip().startswith("4.") or \
                   para_text.strip().startswith("5."):
                    add_body(doc, para_text.strip())
                elif para_text.strip().startswith("☑️"):
                    p = add_body(doc, para_text.strip())
                    for run in p.runs:
                        run.bold = True
                elif para_text.strip().startswith("⚠️"):
                    p = add_body(doc, para_text.strip())
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
                else:
                    add_body(doc, para_text.strip())

    # ── CHECKLIST ──
    add_heading_styled(doc, "종합 체크리스트", 1)
    add_table(doc,
        ["#", "완료 조건", "담당", "기한"],
        [
            ["1", "투자자 주주간 계약서 서명 완료", "CEO+투자자", "7월 10일"],
            ["2", "FBI 범죄증명서 Apostille 완료", "투자자", "7월 10일"],
            ["3", "사무실 임대차계약 + 확정일자", "CEO", "7월 20일"],
            ["4", "정관 공증 완료", "CEO", "7월 23일"],
            ["5", "외국인투자신고 접수 완료", "CEO(투자자 위임)", "7월 25일"],
            ["6", "자본금 1억원 전액 입금", "투자자+CEO", "7월 28일"],
            ["7", "외국인투자기업 등록증 수령", "CEO", "7월 30일"],
            ["8", "법인설립등기 완료", "CEO+법무사", "8월 6일"],
            ["9", "사업자등록증 수령", "CEO", "8월 11일"],
            ["10", "영업보증보험 증권 발급", "CEO", "8월 22일"],
            ["11", "국내외여행업 등록증 수령", "CEO", "9월 10일"],
            ["12", "D-8 비자 신청 서류 송부", "CEO", "9월 15일"],
            ["13", "영업 개시", "CEO", "9월 중순"],
        ]
    )

    add_body(doc, "\n---\n본 문서는 2026년 6월 기준 정보로 작성되었습니다. 실제 진행 시 법무사 및 세무사와 상담하십시오.")

    path2 = OUTDIR + "Resonate_Club_CEO_F4_Action_Plan_KO.docx"
    doc.save(path2)
    print(f"[OK] Document 2 saved: {path2}")
    return path2


# ══════════════════════════════════════════════════════════════════════
# DOCUMENT 3: US Partner (Vendor Manager) Brief (Korean)
# ══════════════════════════════════════════════════════════════════════

def create_doc3_partner_brief():
    doc = Document()
    set_narrow_margins(doc)

    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("리조네이트 클럽 (Resonate Club)\n미국 현지 파트너사 사업계획 및 타임라인")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1a, 0x34, 0x50)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle.add_run("한국→미국 프리미엄 골프 투어 · 현지 벤더 관리 · 공동 운영")
    run2.font.size = Pt(11)
    run2.font.color.rgb = RGBColor(0xd4, 0xaf, 0x37)

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f"작성일: 2026년 6월 11일\n대표이사: 이희성 | 브랜드: 결 투어 (Resonate Tour)\n웹사이트: resonateclub.com / resonatetour.com").font.size = Pt(10)

    doc.add_page_break()

    # ── 1. 사업 개요 ──
    add_heading_styled(doc, "1. 사업 개요", 1)
    add_body(doc,
        "리조네이트 클럽 주식회사(Resonate Club Inc.)는 부산 해운대구에 본사를 둔 프리미엄 여행 기획사로, "
        "'결 투어(Resonate Tour)' 브랜드로 한국인 대상 미국 프리미엄 골프 투어를 주력으로 운영합니다.\n\n"
        "핵심 사업 모델:\n"
        "• 한국인 고소득 골퍼 대상 → 캘리포니아 프리미엄 골프 투어\n"
        "• 올인클루시브 패키지 (항공·숙박·골프·식사·교통·트립릴)\n"
        "• 소규모 프라이빗 그룹 (2~8인) 맞춤형 운영\n"
        "• 원가 +40% 마크업 투명 가격 정책 (고객에게는 원가 비공개)\n"
        "• 추후 미국인→한국 인바운드 럭셔리 투어로 확장\n\n"
        "파트너사의 역할:\n"
        "• 미국 현지에 상주하며 모든 현지 벤더(골프장, 호텔, 교통, 식당 등) 관리\n"
        "• 신규 골프 코스 발굴 및 계약\n"
        "• 투어 운영 중 현지 상황 대응 (실시간 문제 해결)\n"
        "• 고객 만족도 관리 및 피드백 수집\n"
        "• 현지 마케팅 협력 (인스타그램 콘텐츠 등)"
    )

    # ── 2. 법인 설립 타임라인 ──
    add_heading_styled(doc, "2. 법인 설립 타임라인 (파트너사 참고용)", 1)
    add_table(doc,
        ["시기", "이벤트", "의미"],
        [
            ["2026년 7월 17일", "CEO 한국 입국 (F-4 비자)", "법인 설립 시작"],
            ["2026년 7월 말", "외국인투자신고 + 자본금 입금", "1억원 자본금 확보"],
            ["2026년 8월 초", "법인설립등기 완료", "리조네이트 클럽 주식회사 공식 출범"],
            ["2026년 8월 중순", "사업자등록 완료", "세금계산서 발행 가능"],
            ["2026년 9월 초", "국내외여행업 등록 완료", "여행 상품 판매 가능"],
            ["2026년 9월 중순", "영업보증보험 가입 완료", "법적 요건 100% 충족, 영업 개시"],
            ["2026년 10월", "첫 투어 팀 출발 (목표)", "실전 운영 시작"],
        ]
    )

    # ── 3. 패키지 구성 ──
    add_heading_styled(doc, "3. 주요 여행 패키지 구성", 1)

    add_heading_styled(doc, "3.1 한국→미국 골프 투어 (주력)", 2)
    add_table(doc,
        ["패키지명", "기간", "라운드", "판매가 (USD/인)", "주요 코스"],
        [
            ["Trump Premier", "6박7일", "4라운드", "$9,752", "Trump National LA, Sandpiper GC, Hidden Valley GC, The Crossings"],
            ["Trump Elite", "5박6일", "3라운드", "$7,741", "Trump National LA, Sandpiper GC, Hidden Valley GC"],
        ]
    )

    add_heading_styled(doc, "3.2 향후 확장: 미국→한국 인바운드", 2)
    add_table(doc,
        ["패키지명", "기간", "판매가 (USD/인)", "주요 콘텐츠"],
        [
            ["Korea Golf Premier", "11박12일", "$27,300", "나인브릿지·사우스케이프·잭니클라우스GC"],
            ["Korea Discovery", "12박13일", "$27,580", "K컬처 체험 + 골프 + 미쉐린 다이닝"],
        ]
    )

    add_heading_styled(doc, "3.3 VIP 이벤트 (기획 중)", 2)
    add_table(doc,
        ["이벤트명", "일정", "가격", "특징"],
        [
            ["Moving Day VVIP Invitational", "2026년 6월 3-8일", "$14,400/인", "8인 한정, 호스트 Aimee Cho"],
        ]
    )

    # ── 4. 파트너사 역할 ──
    add_heading_styled(doc, "4. 파트너사 세부 역할 및 책임", 1)

    roles = [
        ("4.1 골프 코스 벤더 관리",
         "• 캘리포니아 내 프리미엄 골프 코스 발굴 및 계약\n"
         "  - 현재 계약/협의 중: Trump National LA, Sandpiper GC, Hidden Valley GC, The Crossings, Alisal River Course\n"
         "  - 신규 발굴 목표: Pebble Beach, Torrey Pines, Pelican Hill 등\n"
         "• 티타임 확보: 시즌별·요일별 안정적 티타임 블록 계약\n"
         "• 그린피 협상: 단체 할인(4인+), 리플레이 할인, 시즌 특가 등\n"
         "• 코스 상태 모니터링: 에어레이션 일정, 잔디 상태, 공사 정보 사전 파악"),
        ("4.2 호텔 벤더 관리",
         "• 5성급 호텔 객실 블록 계약\n"
         "  - 주요 협력 호텔: Beverly Wilshire, Terranea Resort, Waldorf Astoria Beverly Hills\n"
         "  - 시즌별 단가 협상, 조식 포함 여부, 취소/변경 조건\n"
         "• 고객 등급별 호텔 옵션 다양화 (5성급 + 4성급 대안)\n"
         "• 체크인/체크아웃 원활한 진행 관리"),
        ("4.3 교통·물류 관리",
         "• 공항 픽업/샌딩 서비스: ICN↔LAX 연계\n"
         "• 투어 전용 차량: 12인승 프리미엄 밴 + 전담 드라이버\n"
         "• 코스 간 이동 동선 최적화 (로드매니저 역할)\n"
         "• 헬리콥터/요트 등 VIP 특별 이동수단 수배\n"
         "• 골프백 운송 관리 (항공 수하물 + 지상 이동)"),
        ("4.4 식음료·레스토랑 관리",
         "• 미쉐린 레스토랑 예약 및 코스 메뉴 조율\n"
         "• 한식 옵션이 있는 레스토랑 리스트 유지\n"
         "• 코스 내 식사(턴하우스/클럽하우스) 사전 주문\n"
         "• 고객 식이 제한(알러지 등) 대응\n"
         "• LA 한인타운 식당과의 제휴 (가이드+고객 식사)"),
        ("4.5 현지 가이드 관리",
         "• 한국어·영어 능통 가이드 섭외 및 관리\n"
         "• 가이드 교육: 코스 정보, 고객 응대 매뉴얼\n"
         "• 투어별 가이드 배정 및 일정 조율\n"
         "• 비상 상황 발생 시 의사소통 채널 유지"),
        ("4.6 마케팅·콘텐츠 협력",
         "• 골프 코스 촬영 (프로모션용 사진·영상)\n"
         "• 인스타그램/유튜브 콘텐츠용 현지 현장 영상 제공\n"
         "• 시네마틱 트립 릴 촬영 지원\n"
         "• 현지 골프 이벤트·토너먼트 정보 공유"),
    ]

    for title_text, body_text in roles:
        add_heading_styled(doc, title_text, 2)
        for line in body_text.split("\n"):
            if line.strip():
                add_body(doc, line.strip())

    # ── 5. 커뮤니케이션 ──
    add_heading_styled(doc, "5. 커뮤니케이션 및 업무 체계", 1)
    add_table(doc,
        ["항목", "내용"],
        [
            ["주요 연락 채널", "카카오톡(실시간), 이메일(공식 문서), Google Drive(문서 공유)"],
            ["정기 미팅", "주 1회 화상 미팅 (Zoom/Google Meet) — 투어 일정 조율, 이슈 점검"],
            ["긴급 연락", "24시간 카카오톡 + 한국·미국 전화 번호 상호 공유"],
            ["문서 공유", "Google Drive: 벤더 계약서, 가격표, 투어 일정표, 고객 정보(암호화)"],
            ["보고 체계", "월간: 벤더 계약 현황 / 분기: 재무 정산 / 수시: 신규 코스·호텔 제안"],
            ["의사 결정", "일상적 결정: 파트너사 자율 / 3천 달러 이상 지출: CEO 승인 / 전략적 결정: 협의"],
        ]
    )

    # ── 6. 재무 ──
    add_heading_styled(doc, "6. 재무 구조 및 수익 배분", 1)

    add_heading_styled(doc, "6.1 비용 구조 (투어 1팀 기준, 4인)", 2)
    add_table(doc,
        ["비용 항목", "추정 원가 (USD)", "비고"],
        [
            ["항공권 (ICN↔LAX 왕복)", "$6,000–8,000", "4인 합계, 이코노미 기준"],
            ["골프 그린피 (4라운드)", "$5,000–8,000", "코스 구성에 따라"],
            ["호텔 숙박 (6박)", "$6,000–12,000", "5성급, 2객실"],
            ["전용 밴 + 드라이버", "$3,000–4,500", "7일 전일정"],
            ["가이드 비용", "$2,000–3,000", "한국어 가이드 1인"],
            ["식사 (전체)", "$2,500–4,000", "레스토랑 + 코스 내 식사"],
            ["기타 (트립릴 등)", "$500–1,000", "촬영·편집·기념품"],
            ["총 원가", "$25,000–40,500", ""],
            ["판매가 (4인)", "$35,000–56,700", "원가+40% 마크업"],
            ["마진 (4인)", "$10,000–16,200", "28–40% 마진율"],
        ]
    )

    add_heading_styled(doc, "6.2 수익 배분 모델 (제안)", 2)
    add_body(doc,
        "파트너사에는 다음과 같은 수익 배분 구조를 제안합니다:\n\n"
        "1. 기본 운영 수수료: 투어당 순이익의 20~30%\n"
        "   - 골프장·호텔·교통·식당 벤더 관리 보상\n"
        "   - 실제 투어 진행 모니터링 및 문제 해결\n\n"
        "2. 신규 벤더 발굴 인센티브: 첫 계약 성사 시 원가 절감액의 30%\n"
        "   - 기존 벤더 대비 낮은 가격이나 더 좋은 조건 확보 시\n\n"
        "3. 시즌 성과 보너스: 연간 목표 초과 달성 시\n"
        "   - 연 12팀 초과 시 팀당 추가 보너스\n\n"
        "4. 인바운드(미국→한국) 확장 시: 인바운드 전담 파트너십 재계약\n"
        "   - 현지 미국인 대상 마케팅 + 모집 대행 수수료 추가\n\n"
        "※ 구체적인 비율과 조건은 상호 협의하여 결정합니다."
    )

    # ── 7. FY2026-27 목표 ──
    add_heading_styled(doc, "7. 연간 운영 목표 (2026년 9월 ~ 2027년 8월)", 1)
    add_table(doc,
        ["구분", "Year 1 목표", "월 평균", "비고"],
        [
            ["투어 팀 수", "12~18팀", "1~1.5팀/월", "보수적 시나리오"],
            ["팀당 평균 인원", "4명", "—", "소규모 프라이빗"],
            ["연간 매출", "약 6억~9억원", "—", "$440K~$660K"],
            ["연간 순이익", "약 1.5억~2.5억원", "—", "마진율 25~35%"],
            ["신규 코스 계약", "3~5곳", "—", "캘리포니아 중심"],
            ["신규 호텔 계약", "2~3곳", "—", "5성급 위주"],
            ["SNS 팔로워", "5,000+", "—", "인스타그램 + 유튜브"],
            ["고객 만족도", "4.8/5.0", "—", "리뷰 평점 목표"],
        ]
    )

    # ── 8. 준비 사항 ──
    add_heading_styled(doc, "8. 파트너사 즉시 준비 사항 (2026년 7월~8월)", 1)
    add_body(doc,
        "1. 사업자 등록 (미국)\n"
        "   - Sole Proprietorship 또는 LLC 설립\n"
        "   - 파트너사 명의의 계약 체결 및 대금 수취 가능하도록\n\n"
        "2. 벤더 사전 접촉\n"
        "   - 기존 관계 활용하여 2026년 하반기 티타임·객실 사전 확보\n"
        "   - Trump National LA: 가을 시즌 티타임 블록 논의\n"
        "   - 호텔: 9~11월 객실 블록 계약 추진\n\n"
        "3. 현장 답사\n"
        "   - 주요 코스 재방문 → 사진·영상 콘텐츠 확보\n"
        "   - 호텔 객실·레스토랑 실사\n"
        "   - 코스 간 이동 시간 실측\n\n"
        "4. 비상 매뉴얼 작성\n"
        "   - 고객 부상/질병 시 병원 리스트\n"
        "   - 항공편 지연/결항 대응 절차\n"
        "   - 천재지변 대응 플랜\n\n"
        "5. 첫 투어 리허설\n"
        "   - 9월 말~10월 초 첫 투어 전 코스·호텔·식당 최종 확인\n"
        "   - CEO와 합동 리허설 진행"
    )

    # ── 마무리 ──
    add_heading_styled(doc, "9. 맺음말", 1)
    add_body(doc,
        "리조네이트 클럽은 단순한 여행사가 아닌, '한국과 미국을 잇는 프리미엄 경험 설계사'입니다.\n\n"
        "파트너사는 미국 현지에서 이 비전을 실현하는 핵심 축입니다. "
        "한국 본사가 상품 기획·마케팅·고객 관리를 담당하는 동안, "
        "파트너사는 최고의 현지 경험을 제공할 수 있도록 모든 인프라를 관리합니다.\n\n"
        "이 파트너십이 성공하기 위한 핵심 요소:\n"
        "• 투명한 커뮤니케이션 — 문제는 숨기지 않고 즉시 공유\n"
        "• 고객 중심 사고 — '고객에게 잊지 못할 경험을'이라는 공동 미션\n"
        "• 지속적 개선 — 매 투어 후 피드백 반영하여 업그레이드\n"
        "• 상호 신뢰 — 정직한 정산, 공정한 수익 배분\n\n"
        "함께 만들어갈 이 여정에 대한 기대와 확신을 담아, 본 계획서를 제안합니다.\n\n"
        "대표이사 이희성 드림\n"
        "resonatetour.com / resonateclub.com"
    )

    path3 = OUTDIR + "Resonate_Club_US_Partner_Brief_KO.docx"
    doc.save(path3)
    print(f"[OK] Document 3 saved: {path3}")
    return path3


# ══════════════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    p1 = create_doc1_us_investor()
    p2 = create_doc2_ceo_actionplan()
    p3 = create_doc3_partner_brief()
    print(f"\n[DONE] All 3 documents created in {OUTDIR}")
    print(f"  1. {p1}")
    print(f"  2. {p2}")
    print(f"  3. {p3}")
