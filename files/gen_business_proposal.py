#!/usr/bin/env python3
"""Generate consolidated Business Proposal (EN + KO) from latest 5 docs."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTDIR = "/Users/hslee/workspace/resonateclub.github.io/files/"

# ─── CONSTANTS ───
C = {
    'capital_total': '100,000,000',
    'capital_min': '50,000,000',
    'insurance_total': '230,000,000',
    'insurance_base': '30,000,000',
    'insurance_planned': '200,000,000',
    'monthly_ins_prem': '192,000',
    'monthly_liability': '35,000',
    'monthly_total': '1,727,000',
    'annual_total': '20,724,000',
    'total_liquidity': '130,000,000',
    'shares': '2,000',
    'price_per_share': '50,000',
    'd8_validity': '2–3 years initially, renewable up to 5 years',
    'f5_req': 'KRW 300M (~$200K) investment OR 2+ full-time Korean employees',
    'markup': '40%',
    'packages': [
        ('Trump Premier', '6N/7D', '4 rounds', '$9,752', 'Trump National LA, Sandpiper GC, Hidden Valley GC, The Crossings'),
        ('Trump Elite', '5N/6D', '3 rounds', '$7,741', 'Trump National LA, Sandpiper GC, Hidden Valley GC'),
    ],
    'future_packages': [
        ('Korea Golf Premier', '11N/12D', '$27,300', 'Nine Bridges, South Cape, Jack Nicklaus GC'),
        ('Korea Discovery', '12N/13D', '$27,580', 'K-Culture + Golf + Michelin Dining'),
    ],
}

def set_cell_shading(cell, color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:fill'), color); sh.set(qn('w:val'), 'clear')
    tcPr.append(sh)

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        set_cell_shading(c, '1a3450')
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = table.rows[ri+1].cells[ci]; c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(9)
            if ri % 2 == 0: set_cell_shading(c, 'f0f4f8')
    doc.add_paragraph()
    return table

def set_margins(doc):
    for s in doc.sections:
        s.top_margin = Cm(1.5); s.bottom_margin = Cm(1.5)
        s.left_margin = Cm(2.0); s.right_margin = Cm(2.0)

def H(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs: r.font.color.rgb = RGBColor(0x1a, 0x34, 0x50)
    return h

def P(doc, text):
    p = doc.add_paragraph(text); p.style.font.size = Pt(10)
    p.style.paragraph_format.space_after = Pt(6)
    return p

def cover(doc, title_en, title_ko, subtitle, date_line):
    doc.add_paragraph(); doc.add_paragraph()
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(f"{title_en}\n{title_ko}"); r.bold = True
    r.font.size = Pt(20); r.font.color.rgb = RGBColor(0x1a, 0x34, 0x50)
    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = s.add_run(subtitle); r2.font.size = Pt(12); r2.font.color.rgb = RGBColor(0xd4, 0xaf, 0x37)
    doc.add_paragraph()
    i = doc.add_paragraph(); i.alignment = WD_ALIGN_PARAGRAPH.CENTER
    i.add_run(date_line).font.size = Pt(10)
    doc.add_page_break()


# ═══════════════════════════════════════════════════════
# ENGLISH PROPOSAL
# ═══════════════════════════════════════════════════════
def create_en():
    doc = Document(); set_margins(doc)
    cover(doc, "RESONATE CLUB INC.", "리조네이트 클럽 주식회사",
          "Premium Korea↔USA Travel | Business Proposal",
          "Date: June 11, 2026\nInvestor: U.S. Citizen/Entity (100% Equity)\nCEO: Hee Sung Lee (F-4 Visa)\nBrand: Resonate Tour (결 투어)")

    # 1. EXECUTIVE SUMMARY
    H(doc, "1. Executive Summary", 1)
    P(doc,
        f"Resonate Club Inc. is a premium travel planning corporation to be established in Haeundae-gu, "
        f"Busan, Republic of Korea, under the Foreign Investment Promotion Act (FIPA). Operating under "
        f"the brand 'Resonate Tour (결 투어),' the company specializes in all-inclusive, small-group "
        f"premium golf tours from Korea to California.\n\n"
        f"The company will be registered as a Domestic & International Travel Agency with Planned Travel "
        f"Operator status (국내외여행업 / 기획여행사), with a total paid-in capital of KRW {C['capital_total']} "
        f"(~USD $74,100), fully funded by a U.S. investor holding 100% equity. Hee Sung Lee, an F-4 "
        f"Overseas Korean Visa holder, will serve as non-shareholder Representative Director (CEO).\n\n"
        f"Key financials: Monthly fixed costs of KRW {C['monthly_total']}, guarantee insurance of KRW "
        f"{C['insurance_total']} (Base KRW {C['insurance_base']} + Planned Travel KRW {C['insurance_planned']}), "
        f"and a cost-plus-{C['markup']} markup model targeting 12–18 premium golf tour groups in Year 1."
    )

    # 2. COMPANY OVERVIEW
    H(doc, "2. Company Overview", 1)
    add_table(doc,
        ["Item", "Detail"],
        [
            ["Legal Entity", "Resonate Club Inc. — Foreign-Invested Joint-Stock Company (주식회사)"],
            ["Brand", "Resonate Tour (결 투어)"],
            ["Business Type", "Domestic & International Travel Agency — Planned Travel Operator"],
            ["Registered Office", "Haeundae-gu, Busan, Republic of Korea"],
            ["Paid-in Capital", f"KRW {C['capital_total']} (statutory minimum KRW {C['capital_min']})"],
            ["Shares", f"{C['shares']} common shares at KRW {C['price_per_share']}/share"],
            ["Ownership", "U.S. Investor: 100% (sole shareholder)"],
            ["CEO", "Hee Sung Lee — F-4 Visa, non-shareholder Representative Director (salaried)"],
            ["Planned Establishment", "August 2026 — Operations from late September 2026"],
            ["Web", "resonatetour.com / resonateclub.com"],
        ]
    )

    # 3. BUSINESS MODEL
    H(doc, "3. Business Model & Services", 1)

    H(doc, "3.1 Primary Service: Korea→USA Premium Golf Tours", 2)
    P(doc,
        f"Target: High-net-worth Korean golfers (annual income KRW 100M+, ages 35–60)\n"
        f"Operating model: Small private groups (2–8 persons), all-inclusive, fully customized\n"
        f"Pricing: Cost + {C['markup']} transparent markup (cost never disclosed to clients)\n"
        f"Included: Round-trip flights, 5-star hotels, premium golf courses, private van + driver, "
        f"Korean-speaking guide, all meals, cinematic trip reel"
    )
    add_table(doc,
        ["Package", "Duration", "Rounds", "Price (USD/person)", "Key Courses"],
        [list(p) for p in C['packages']]
    )

    H(doc, "3.2 Future Expansion: USA→Korea Inbound Luxury Tours", 2)
    P(doc, "After upgrading to a Comprehensive Travel Agency (종합여행업) license, the company will expand to serve American luxury travelers visiting Korea for K-culture, premium golf, and Michelin dining experiences.")
    add_table(doc,
        ["Package", "Duration", "Price (USD/person)", "Highlights"],
        [list(p) for p in C['future_packages']]
    )

    H(doc, "3.3 VIP Events", 2)
    P(doc, "Curated, ultra-exclusive events such as the Moving Day VVIP Invitational (8 guests only, $14,400/person) hosted by Aimee Cho, offering access and experiences unavailable through mass-market channels.")

    # 4. CORPORATE STRUCTURE
    H(doc, "4. Corporate Structure & Governance", 1)
    add_table(doc,
        ["Role", "Person", "Status", "Key Responsibilities"],
        [
            ["Sole Shareholder (100%)", "U.S. Investor", "D-8-1 Corporate Investment Visa", "Capital provision, strategic oversight, board-level decisions"],
            ["Representative Director (CEO)", "Hee Sung Lee", "F-4 Overseas Korean Visa\nNon-shareholder, salaried", "Day-to-day operations, licensing, banking, contracts, tax filings, government relations"],
            ["Director / Auditor", "U.S. Investor", "Registered on corporate registry", "Legal compliance, major transaction approval"],
            ["U.S. Partner\n(Vendor Manager)", "TBD", "Based in California", "Local vendor management, course scouting, on-site tour operations, content creation"],
        ]
    )
    P(doc,
        "The CEO operates under a comprehensive Power of Attorney from the investor, enabling full "
        "autonomy for daily operations. Major decisions (capital changes, M&A, dividend declarations) "
        "require board resolution with investor approval."
    )

    # 5. ESTABLISHMENT TIMELINE
    H(doc, "5. Establishment Timeline", 1)
    add_table(doc,
        ["Phase", "Dates", "Key Actions", "Outcome"],
        [
            ["Pre-Setup", "Now – Jul 16", "Shareholder agreement, investor docs, engage Korean professionals", "Ready for CEO arrival"],
            ["STEP 1", "Jul 17–20", "CEO arrival Korea (F-4), Haeundae-gu office lease", "Office secured"],
            ["STEP 2", "Jul 21–23", "Articles of Incorporation + notarization", "Notarized Articles"],
            ["STEP 3", "Jul 24–30", "FDI notification → Capital deposit (KRW 100M)", "FDI registration + bank certificate"],
            ["STEP 4", "Jul 31 – Aug 6", "Corporation registration at Busan District Court", "Legal entity established"],
            ["STEP 5", "Aug 7–11", "Business registration at Haeundae Tax Office", "Tax ID + business certificate"],
            ["STEP 6", "Aug 12 – Sep 10", "Travel Agency License application (Haeundae-gu Office)", "Domestic & Int'l Travel Agency License"],
            ["STEP 7", "Aug 12 – Sep 17", f"Guarantee Insurance (SGI) — KRW {C['insurance_total']} + Liability Insurance", "Full insurance compliance"],
            ["STEP 8", "Sep 15 onward", "Investor D-8 visa application (US Embassy)", "Investor residency secured"],
            ["Operations", "Late Sep 2026", "First tour group departure", "Revenue generation begins"],
        ]
    )

    # 6. FINANCIAL PLAN
    H(doc, "6. Financial Plan", 1)

    H(doc, "6.1 Capital & Initial Costs", 2)
    add_table(doc,
        ["Category", "Amount (KRW)", "USD (~1,350 rate)", "Note"],
        [
            ["Paid-in Capital", C['capital_total'], "$74,100", f"Exceeds statutory minimum (KRW {C['capital_min']})"],
            ["Initial Setup (registration, notary, office deposit)", "~5,000,000", "~$3,700", "One-time costs"],
            ["1-Year Operating Reserve", "~25,000,000", "~$18,500", "Monthly fixed × 12 + contingency"],
            ["TOTAL Recommended", C['total_liquidity'], "~$96,300", "Covers first year of operations"],
        ]
    )

    H(doc, "6.2 Monthly Fixed Costs", 2)
    add_table(doc,
        ["Item", "Monthly (KRW)", "Annual (KRW)", "Note"],
        [
            ["Office Rent (Haeundae-gu)", "800,000", "9,600,000", "Small office, premium location"],
            ["Guarantee Insurance (SGI)", C['monthly_ins_prem'], "2,300,000", f"Base {C['insurance_base']} + Planned Travel {C['insurance_planned']} = {C['insurance_total']} bond; ~1% premium"],
            ["Liability Insurance", C['monthly_liability'], "420,000", "Mandatory under Tourism Promotion Act"],
            ["Accounting & Tax Services", "300,000", "3,600,000", "Monthly bookkeeping + annual tax filing"],
            ["Telecom & Office Supplies", "200,000", "2,400,000", "Internet, phone, sundries"],
            ["Miscellaneous", "200,000", "2,400,000", "Bank fees, memberships, contingency"],
            ["Monthly Total", C['monthly_total'], C['annual_total'], ""],
        ]
    )

    H(doc, "6.3 Revenue Model", 2)
    P(doc,
        f"Core model: Cost-Plus {C['markup']} Markup\n\n"
        f"Year 1 Conservative Target (Sep 2026 – Aug 2027):\n"
        f"• 12–18 tour groups × avg 4 persons × ~$2,500 profit/person\n"
        f"• Gross profit: ~$120,000–180,000\n"
        f"• After operating costs: ~KRW 130M–200M (~$96K–148K) net profit\n\n"
        f"Year 2 (full calendar, 24+ groups): ~KRW 300M+ (~$222K+) net profit\n\n"
        f"Profit repatriation via dividends (10–15% Korean withholding, creditable against U.S. tax via Form 1116)."
    )

    H(doc, "6.4 Guarantee Insurance Breakdown", 2)
    P(doc,
        f"As a Planned Travel Operator (기획여행사), the company must carry guarantee insurance "
        f"protecting customer prepayments. The Tourism Promotion Act Enforcement Rules Schedule 3 require:\n"
        f"  • Base bond (국내외여행업): KRW {C['insurance_base']}\n"
        f"  • Planned Travel surcharge (기획여행 추가): KRW {C['insurance_planned']}\n"
        f"  • TOTAL: KRW {C['insurance_total']}\n\n"
        f"This is an insurance PREMIUM (annual ~KRW 2.3M at ~1% rate), not a deposit. "
        f"Additionally, mandatory Liability Insurance (배상책임보험) covers bodily/property damage during tours."
    )

    # 7. LICENSE & COMPLIANCE
    H(doc, "7. Licensing & Regulatory Compliance", 1)
    add_table(doc,
        ["License/Permit", "Legal Basis", "Authority", "Est. Duration"],
        [
            ["Corporation Registration", "Commercial Act", "Busan District Court (East)", "5–7 days"],
            ["Business Registration", "VAT Act", "Haeundae Tax Office", "3–5 days"],
            ["Travel Agency License", "Tourism Promotion Act Art. 4", "Haeundae-gu Office", "2–4 weeks"],
            ["Guarantee Insurance (SGI)", "Tourism Act Enforcement Rules Sch. 3", "Seoul Guarantee Insurance", "5–7 days"],
            ["Liability Insurance", "Tourism Promotion Act", "Private insurer", "3–5 days"],
            ["FDI Registration", "FIPA", "KOTRA / Foreign Exchange Bank", "1–2 days"],
            ["Tour Escort Registration", "Tourism Promotion Act Art. 38", "Korea Tourism Organization", "Hire qualified"],
        ]
    )

    # 8. D-8 VISA & TAX BENEFITS
    H(doc, "8. Investor Visa & Tax Benefits", 1)
    H(doc, "8.1 D-8-1 Corporate Investment Visa", 2)
    add_table(doc,
        ["Benefit", "Details"],
        [
            ["Visa Type", "D-8-1 (Corporate Investment)"],
            ["Validity", C['d8_validity'] + ", renewable indefinitely"],
            ["Family", "Spouse + minor children eligible for F-3 dependent visas"],
            ["Healthcare", "Full NHIS (National Health Insurance) coverage"],
            ["PR Pathway", "F-5 Permanent Residency: " + C['f5_req']],
        ]
    )

    H(doc, "8.2 FIPA Tax Incentives", 2)
    add_table(doc,
        ["Incentive", "Benefit", "Duration"],
        [
            ["Corporate Income Tax Reduction", "100% exemption years 1–5, 50% years 6–7", "7 years"],
            ["Acquisition Tax Exemption", "100% on real estate, vehicles, equipment", "15 years"],
            ["Property Tax Reduction", "100% years 1–5, 50% years 6–15", "15 years"],
            ["Customs Duty Exemption", "100% on imported capital goods", "Within 3 years of FDI"],
            ["Dividend Withholding (US-Korea Treaty)", "15% standard / 10% if held via corporate entity with ≥10% ownership", "Ongoing"],
        ]
    )

    # 9. TEAM
    H(doc, "9. Team & Key Partners", 1)
    P(doc,
        "U.S. INVESTOR (100% Shareholder): Provides capital, strategic direction, and board-level oversight. "
        "Holds D-8-1 visa for Korean residency. Engages U.S. cross-border CPA for Form 5471, FBAR, and "
        "FATCA compliance.\n\n"
        "HEE SUNG LEE (CEO, F-4 Visa): 20+ years U.S. residency, hands-on golf industry experience at "
        "American Golf Inc. / Recreation Park 18 Golf Course, extensive California/Hawaii course network, "
        "bilingual Korean-English. Manages all Korean operations: licensing, banking, tax, marketing, "
        "client relationships.\n\n"
        "U.S. PARTNER (California-Based Vendor Manager): Manages all local vendors — golf courses, "
        "hotels, transport, restaurants. Scouts new courses, negotiates contracts, handles real-time "
        "tour operations, creates social media content. Revenue share: 20–30% of net profit per tour.\n\n"
        "KOREAN CERTIFIED TAX ACCOUNTANT (세무사): Corporate tax filing, VAT returns, FIPA incentive "
        "applications, withholding tax on dividends. KRW 300,000–500,000/month.\n\n"
        "U.S. CROSS-BORDER CPA: Form 5471 (foreign corporation reporting), FBAR, Form 1116 (foreign "
        "tax credit), Form 8938 (FATCA), PFIC analysis. $2,500–5,000 initial / $1,500–3,000/year."
    )

    # 10. INVESTMENT CASE
    H(doc, "10. Investment Case", 1)
    P(doc,
        "Why invest in Resonate Club Inc.?\n\n"
        "1. STRUCTURAL MARKET GROWTH: Korean golf population 6.24M, overseas golf travel growing at "
        "12.5% CAGR. U.S. is the #2 destination for Korean outbound golf.\n\n"
        "2. PREMIUM NICHE: Large agencies serve mass market; nobody owns the premium, customized, "
        "small-group California golf segment for Korean high-net-worth individuals.\n\n"
        "3. UNFAIR ADVANTAGE: CEO's 20+ years U.S. golf industry experience, direct course relationships, "
        "bilingual/bicultural capability — impossible for Korean-only competitors to replicate.\n\n"
        "4. CAPITAL-EFFICIENT: Low fixed cost base (KRW 1.73M/month), asset-light model, 40% markup "
        "provides predictable margins. Breakeven at ~5 tour groups per year.\n\n"
        "5. SCALABLE: Model expands to USA→Korea inbound (K-culture) and VIP events without "
        "proportional cost increase.\n\n"
        "6. TAX-ADVANTAGED: FIPA corporate tax holiday (years 1–5), customs duty exemption on vehicles "
        "and equipment, US-Korea tax treaty reduces dividend withholding.\n\n"
        "7. CLEAN GOVERNANCE: 100% investor ownership, professional CEO, transparent cost-plus pricing, "
        "single-class share structure."
    )

    # 11. RISK
    H(doc, "11. Risk Factors", 1)
    add_table(doc,
        ["Risk", "Level", "Mitigation"],
        [
            ["Regulatory change (capital increase)", "Medium", "Investor-backed capital exceeds requirements; legal counsel retained"],
            ["SGI insurance denial", "Medium", "Strong business plan; Korea Credit Guarantee Fund backup route"],
            ["D-8 visa delay", "Low", "Early application; 5–10 business days typical processing"],
            ["FX fluctuation (KRW/USD)", "Low", "Revenue in USD, costs split — natural hedge"],
            ["Slow Season 1 adoption", "Medium", "Conservative targets; 5 groups = breakeven"],
            ["US tax compliance penalties", "High", "Cross-border CPA engaged pre-investment; Form 5471 filed annually"],
            ["PFIC classification", "Low", "Active operating business; CPA opinion letter secured"],
        ]
    )

    P(doc, "\n---\nThis business proposal synthesizes the latest corporate establishment plans as of June 11, 2026. All figures are based on the Tourism Promotion Act and its Enforcement Rules as currently in effect.")
    path = OUTDIR + "Resonate_Club_Business_Proposal_EN.docx"
    doc.save(path)
    print(f"[OK] EN: {path}")
    return path


# ═══════════════════════════════════════════════════════
# KOREAN PROPOSAL
# ═══════════════════════════════════════════════════════
def create_ko():
    doc = Document(); set_margins(doc)
    cover(doc, "RESONATE CLUB INC.", "리조네이트 클럽 주식회사",
          "프리미엄 한미 여행 플랫폼 | 사업제안서",
          "작성일: 2026년 6월 11일\n투자자: 미국 시민권자/법인 (지분 100%)\n대표이사: 이희성 (F-4 재외동포 비자)\n브랜드: 결 투어 (Resonate Tour)")

    H(doc, "1. 사업 개요", 1)
    P(doc,
        f"리조네이트 클럽 주식회사(Resonate Club Inc.)는 부산 해운대구에 설립 예정인 외국인투자기업으로, "
        f"'결 투어(Resonate Tour)' 브랜드로 한국인 대상 미국 프리미엄 골프 투어를 주력 사업으로 합니다.\n\n"
        f"국내외여행업(기획여행사)으로 등록 예정이며, 총 납입자본금 {C['capital_total']}원은 미국 투자자가 "
        f"100% 출자합니다. 대표이사 이희성은 F-4 재외동포 비자 소지자로서 비주주 대표이사로 취임하여 "
        f"한국 내 모든 운영을 총괄합니다.\n\n"
        f"핵심 지표: 월고정비 ₩{C['monthly_total']}, 영업보증보험 {C['insurance_total']}원 "
        f"(기본 {C['insurance_base']}원 + 기획여행 추가 {C['insurance_planned']}원), "
        f"원가+{C['markup']} 마크업 모델, Year 1 목표 12~18팀."
    )

    H(doc, "2. 회사 개요", 1)
    add_table(doc,
        ["항목", "내용"],
        [
            ["법인명", "리조네이트 클럽 주식회사 (Resonate Club Inc.) — 외국인투자기업 주식회사"],
            ["브랜드", "결 투어 (Resonate Tour)"],
            ["사업 유형", "국내외여행업 (기획여행사 / Planned Travel Operator)"],
            ["본점 소재지", "부산광역시 해운대구"],
            ["납입자본금", f"{C['capital_total']}원 (법정 최소 {C['capital_min']}원)"],
            ["발행주식", f"보통주 {C['shares']}주 (1주당 {C['price_per_share']}원)"],
            ["지분 구조", "미국 투자자 100% (단독 주주)"],
            ["대표이사", "이희성 — F-4 비자, 비주주 대표이사 (급여 수령)"],
            ["설립 예정", "2026년 8월 — 9월 말부터 영업 개시"],
            ["웹사이트", "resonatetour.com / resonateclub.com"],
        ]
    )

    H(doc, "3. 사업 모델", 1)
    H(doc, "3.1 주력 사업: 한국→미국 프리미엄 골프 투어", 2)
    P(doc,
        f"대상: 연소득 1억원 이상 한국인 고소득 골퍼 (35~60세)\n"
        f"운영 방식: 소규모 프라이빗 그룹(2~8인), 올인클루시브, 완전 맞춤형\n"
        f"가격 정책: 원가 + {C['markup']} 마크업 (고객에게는 원가 절대 비공개)\n"
        f"포함: 왕복항공, 5성급호텔, 프리미엄골프, 전용밴+기사, 한국어가이드, 전식사, 시네마틱트립릴"
    )
    add_table(doc,
        ["패키지명", "기간", "라운드", "판매가 (USD/인)", "주요 코스"],
        [list(p) for p in C['packages']]
    )

    H(doc, "3.2 향후 확장: 미국→한국 인바운드 럭셔리 투어", 2)
    P(doc, "종합여행업으로 전환 후, K컬처·프리미엄 골프·미쉐린 다이닝을 결합한 미국인 대상 한국 럭셔리 투어로 확장 예정입니다.")
    add_table(doc,
        ["패키지명", "기간", "판매가 (USD/인)", "주요 콘텐츠"],
        [list(p) for p in C['future_packages']]
    )

    H(doc, "3.3 VIP 이벤트", 2)
    P(doc, "Moving Day VVIP Invitational (8인 한정, $14,400/인, 호스트 Aimee Cho) 등 초특급 이벤트를 기획 중입니다.")

    H(doc, "4. 지배구조", 1)
    add_table(doc,
        ["역할", "담당자", "지위", "주요 책임"],
        [
            ["단독 주주 (100%)", "미국 투자자", "D-8-1 기업투자비자", "자본금 출자, 전략적 의사결정, 이사회 승인"],
            ["대표이사 (CEO)", "이희성", "F-4 재외동포 비자\n비주주, 급여 수령", "일상 운영 총괄, 인허가, 금융, 계약, 세무, 정부 관계"],
            ["이사/감사", "미국 투자자", "법인등기부 등재", "법적 준거성 감독, 주요 거래 승인"],
            ["미국 파트너\n(현지 벤더 관리)", "미정", "캘리포니아 거주", "현지 벤더 관리, 코스 발굴, 투어 현장 운영, 콘텐츠 제작"],
        ]
    )
    P(doc, "대표이사는 투자자로부터 포괄적 위임장(POA)을 수여받아 일상 업무를 자율적으로 수행하며, 자본금 변동·M&A·배당 등 주요 결정은 이사회 결의를 통해 투자자 승인을 받습니다.")

    H(doc, "5. 설립 타임라인", 1)
    add_table(doc,
        ["단계", "일정", "주요 내용", "결과물"],
        [
            ["사전 준비", "현재 ~ 7/16", "주주계약서, 투자자 서류, 한국 전문가 섭외", "입국 준비 완료"],
            ["STEP 1", "7/17~20", "CEO 입국(F-4), 해운대구 사무실 계약", "사무실 확보"],
            ["STEP 2", "7/21~23", "정관 작성 및 공증", "공증된 정관"],
            ["STEP 3", "7/24~30", "외국인투자신고 → 자본금 1억원 입금", "FDI 등록증 + 입금증명서"],
            ["STEP 4", "7/31~8/6", "부산지방법원 동부지원 법인설립등기", "법인 설립 완료"],
            ["STEP 5", "8/7~11", "해운대세무서 사업자등록", "사업자등록증"],
            ["STEP 6", "8/12~9/10", "해운대구청 국내외여행업 등록 신청", "여행업 등록증"],
            ["STEP 7", "8/12~9/17", f"SGI 영업보증보험({C['insurance_total']}원) + 배상책임보험", "보험 요건 충족"],
            ["STEP 8", "9/15~", "투자자 D-8 비자 신청 (주한미대사관)", "투자자 체류 자격 확보"],
            ["영업 개시", "9월 말", "첫 투어팀 출발", "매출 발생 시작"],
        ]
    )

    H(doc, "6. 재무 계획", 1)
    H(doc, "6.1 자본금 및 초기 비용", 2)
    add_table(doc,
        ["구분", "금액 (원)", "비고"],
        [
            ["납입자본금", C['capital_total'], f"법정 최소 {C['capital_min']}원 상회"],
            ["초기 설립 비용 (등기·공증·사무실 보증금 등)", "약 5,000,000", "일회성 비용"],
            ["1년분 운영 적립금", "약 25,000,000", f"월고정비 ₩{C['monthly_total']} × 12개월 + 예비비"],
            ["총 권장 유동성", C['total_liquidity'], "첫해 안정적 운영 보장"],
        ]
    )

    H(doc, "6.2 월간 고정비", 2)
    add_table(doc,
        ["항목", "월 (원)", "연 (원)", "비고"],
        [
            ["사무실 월세 (해운대구)", "800,000", "9,600,000", "소형 오피스, 프리미엄 입지"],
            ["영업보증보험 (SGI)", C['monthly_ins_prem'], "2,300,000", f"기본 {C['insurance_base']}원 + 기획여행 {C['insurance_planned']}원 = 총{C['insurance_total']}원, 약 1% 보험료"],
            ["배상책임보험", C['monthly_liability'], "420,000", "관광진흥법상 의무 가입"],
            ["회계·세무 대행", "300,000", "3,600,000", "월 기장 + 연간 법인세 신고"],
            ["통신·사무용품", "200,000", "2,400,000", "인터넷, 전화, 소모품"],
            ["기타 잡비", "200,000", "2,400,000", "은행 수수료, 협회비, 예비비"],
            ["월간 합계", C['monthly_total'], C['annual_total'], ""],
        ]
    )

    H(doc, "6.3 수익 모델", 2)
    P(doc,
        f"핵심 모델: 원가 + {C['markup']} 마크업\n\n"
        f"Year 1 보수적 목표 (2026년 9월 ~ 2027년 8월):\n"
        f"• 12~18팀 × 평균 4인 × 인당 약 $2,500 이익\n"
        f"• 매출총이익: $120,000~180,000\n"
        f"• 고정비 차감 후 순이익: 약 1.3억~2억원\n\n"
        f"Year 2 (연간 24팀+): 순이익 약 3억원+\n\n"
        f"이익 배당 시 한미조세조약에 따라 원천세 10~15% (미국에서 외국납부세액공제 가능)."
    )

    H(doc, "6.4 영업보증보험 상세", 2)
    P(doc,
        f"기획여행사는 고객 선급금을 보호하기 위해 관광진흥법 시행규칙 별표3에 따라 "
        f"다음 보증보험에 가입해야 합니다:\n"
        f"  • 기본 보증(국내외여행업): {C['insurance_base']}원\n"
        f"  • 기획여행 추가 보증: {C['insurance_planned']}원\n"
        f"  • 합계: {C['insurance_total']}원\n\n"
        f"연 보험료는 보증금액의 약 1% 수준(약 230만원). 예치금이 아닌 보험료 방식이므로 "
        f"자본금과 별도로 운영 가능합니다. 배상책임보험(월 {C['monthly_liability']}원)도 의무 가입입니다."
    )

    H(doc, "7. 인허가 사항", 1)
    add_table(doc,
        ["인허가", "근거법령", "발급기관", "소요기간"],
        [
            ["법인설립등기", "상법", "부산지방법원 동부지원", "5~7일"],
            ["사업자등록", "부가가치세법", "해운대세무서", "3~5일"],
            ["국내외여행업 등록", "관광진흥법 제4조", "해운대구청", "2~4주"],
            ["영업보증보험", "관광진흥법 시행규칙 별표3", "서울보증보험(SGI)", "5~7일"],
            ["배상책임보험", "관광진흥법", "민간 보험사", "3~5일"],
            ["외국인투자신고", "외국인투자촉진법", "KOTRA/외환은행", "1~2일"],
            ["국외여행인솔자 등록", "관광진흥법 제38조", "한국관광공사", "자격자 채용"],
        ]
    )

    H(doc, "8. 투자비자 및 세제혜택", 1)
    H(doc, "8.1 D-8-1 기업투자비자", 2)
    add_table(doc,
        ["혜택", "내용"],
        [
            ["비자 유형", "D-8-1 (기업투자)"],
            ["유효기간", C['d8_validity'] + ", 갱신 가능"],
            ["가족", "배우자 및 미성년 자녀 F-3 동반비자"],
            ["의료", "국민건강보험(NHIS) 적용"],
            ["영주권 경로", "F-5 영주권: " + C['f5_req']],
        ]
    )

    H(doc, "8.2 FIPA 세제 혜택", 2)
    add_table(doc,
        ["혜택", "내용", "기간"],
        [
            ["법인세 감면", "1~5년차 100% 면제, 6~7년차 50% 감면", "7년"],
            ["취득세 면제", "사업용 부동산·차량·장비 100% 면제", "15년"],
            ["재산세 감면", "1~5년차 100%, 6~15년차 50% 감면", "15년"],
            ["관세 면제", "수입 자본재 100% 면제", "FDI 등록 후 3년 이내"],
            ["배당 원천세 (한미조세조약)", "일반 15% / 법인 10%+ 지분 시 10%", "지속"],
        ]
    )

    H(doc, "9. 팀 구성", 1)
    P(doc,
        "미국 투자자 (100% 주주): 자본금 출자, 전략적 방향 제시, 이사회 승인. D-8-1 비자로 한국 체류 가능. "
        "미국 세무사를 통해 Form 5471, FBAR, FATCA 신고.\n\n"
        "이희성 대표이사 (F-4 비자): 미국 20년+ 거주, American Golf Inc. / Recreation Park 18 GC 근무 경력, "
        "캘리포니아·하와이 골프 코스 네트워크, 한영 bilingual. 한국 내 모든 운영 총괄.\n\n"
        "미국 파트너 (캘리포니아 현지): 골프장·호텔·교통·식당 벤더 관리, 신규 코스 발굴, 현장 투어 운영, "
        "SNS 콘텐츠 제작. 수익 배분: 투어당 순이익의 20~30%.\n\n"
        "한국 세무사: 법인세·부가세 신고, FIPA 세제혜택 신청, 배당원천징수. 월 30~50만원.\n\n"
        "미국 국제세무 CPA: Form 5471, FBAR, Form 1116, Form 8938, PFIC 분석. 초기 $2,500~5,000 / 연 $1,500~3,000."
    )

    H(doc, "10. 투자 포인트", 1)
    P(doc,
        "리조네이트 클럽에 투자해야 하는 이유:\n\n"
        "1. 구조적 시장 성장: 한국 골프인구 624만명, 해외골프여행 연 12.5% 성장. 미국은 한국인 골프여행 2위 목적지.\n\n"
        "2. 프리미엄 틈새: 대형 여행사는 대중시장만 커버. 한국 고소득층 대상 캘리포니아 프리미엄 골프 세그먼트는 무주공산.\n\n"
        "3. 넘을 수 없는 진입장벽: CEO의 20년+ 미국 골프업계 경력, 현지 코스 직접 관계, bilingual·bicultural 역량 — 한국 내 경쟁사가 복제 불가.\n\n"
        "4. 자본 효율성: 월고정비 173만원, 자산 경량화 모델, 40% 마크업으로 예측 가능한 수익. 연 5팀만으로 손익분기점 달성.\n\n"
        "5. 확장성: 한국→미국 모델을 미국→한국 인바운드(K컬처)로 확장 가능. VIP 이벤트 추가 수익원.\n\n"
        "6. 세제 혜택: FIPA 법인세 5년 면제, 관세 면제, 한미조세조약 배당 원천세 인하.\n\n"
        "7. 깨끗한 지배구조: 투자자 100% 지분, 전문 CEO 경영, 단일 클래스 주식, 투명한 원가+마크업 가격 정책."
    )

    H(doc, "11. 위험 요소", 1)
    add_table(doc,
        ["위험", "수준", "완화 방안"],
        [
            ["법령 개정 (자본금 인상)", "중", "투자자 자본 여유 충분, 법무사 자문"],
            ["SGI 보증보험 거절", "중", "탄탄한 사업계획서, 신용보증기금 대체 경로"],
            ["D-8 비자 지연", "하", "조기 신청, 미대사관 5~10영업일 소요"],
            ["환율 변동 (원/달러)", "하", "매출 USD, 비용 분산 — 자연 헤지"],
            ["시즌 1 저조", "중", "보수적 목표, 5팀 손익분기점"],
            ["미국 세무 컴플라이언스 위반", "상", "투자 전 CPA 선임, Form 5471 매년 신고"],
            ["PFIC 분류 위험", "하", "적극적 영업회사, CPA 의견서 확보"],
        ]
    )

    P(doc, "\n---\n본 사업제안서는 2026년 6월 11일 기준 최신 법인설립 계획을 종합한 것입니다. 모든 수치는 현행 관광진흥법 및 동법 시행규칙에 근거합니다.")
    path = OUTDIR + "Resonate_Club_Business_Proposal_KO.docx"
    doc.save(path)
    print(f"[OK] KO: {path}")
    return path


if __name__ == "__main__":
    create_en()
    create_ko()
    print("\n[DONE] 2 business proposals generated.")
