#!/usr/bin/env python3
"""Generate Doc 2 (KO+EN) and Doc 3 (KO+EN) with consistent numbers from Doc 1."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

OUTDIR = "/Users/hslee/workspace/resonateclub.github.io/files/"

# ─── SHARED CONSTANTS (sourced from Doc 1 final) ───
CAPITAL_TOTAL = "100,000,000"
CAPITAL_INVESTOR = "100,000,000"  # 100% ownership
CAPITAL_CEO = "0 (non-shareholder CEO)"
CAPITAL_MIN = "50,000,000"
INSURANCE_BASE = "30,000,000"
INSURANCE_PLANNED = "200,000,000"
INSURANCE_TOTAL = "230,000,000"
MONTHLY_INSURANCE_PREM = "192,000"
ANNUAL_INSURANCE_PREM = "2,300,000"
MONTHLY_LIABILITY = "35,000"
ANNUAL_LIABILITY = "420,000"
MONTHLY_OFFICE = "800,000"
MONTHLY_ACCOUNTING = "300,000"
MONTHLY_TELECOM = "200,000"
MONTHLY_MISC = "200,000"
MONTHLY_TOTAL = "1,727,000"
ANNUAL_TOTAL = "20,724,000"
TOTAL_LIQUIDITY = "130,000,000"
SHARES = "2,000"
PRICE_PER_SHARE = "50,000"
D8_VALIDITY = "2–3 years initially, renewable up to 5 years"
F5_REQ = "KRW 300M (~$200K) investment OR 2+ full-time Korean employees"
# Ownership strings
OWNERSHIP_KO = "투자자 100% (CEO는 비주주 대표이사)"
OWNERSHIP_EN = "Investor 100% (CEO is non-shareholder Representative Director)"
EQUITY_KO = "투자자: 미국인, 자본금 {CAPITAL_INVESTOR}원 전액 출자 (지분 100%), D-8-1 비자 취득 예정"
EQUITY_EN = "Investor: U.S. citizen/entity, sole shareholder (100% equity), D-8-1 visa"
CEO_KO = "대표이사: 이희성, F-4 비자, 비주주 대표이사 (급여 수령)"
CEO_EN = "CEO: Hee Sung Lee, F-4 visa, non-shareholder Representative Director (salaried)"

def set_cell_shading(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    tcPr.append(shading)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '1a3450')
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r+1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if r % 2 == 0:
                set_cell_shading(cell, 'f0f4f8')
    doc.add_paragraph()
    return table

def set_narrow_margins(doc):
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x34, 0x50)
    return h

def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.style.font.size = Pt(10)
    p.style.paragraph_format.space_after = Pt(6)
    return p

def add_steps(doc, steps):
    for title, body in steps:
        add_heading_styled(doc, title, 1)
        for para_text in body.split("\n\n"):
            if para_text.strip():
                pt = para_text.strip()
                if pt.startswith(("1.", "2.", "3.", "4.", "5.")):
                    add_body(doc, pt)
                elif pt.startswith("☑"):
                    p = add_body(doc, pt)
                    for run in p.runs:
                        run.bold = True
                elif pt.startswith("⚠"):
                    p = add_body(doc, pt)
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
                else:
                    add_body(doc, pt)


# ══════════════════════════════════════════════════════════════
# DOC 2 — KOREAN
# ══════════════════════════════════════════════════════════════
def create_doc2_ko():
    doc = Document()
    set_narrow_margins(doc)

    doc.add_paragraph(); doc.add_paragraph()
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("리조네이트 클럽 주식회사\nF-4 대표이사 실행계획")
    r.bold = True; r.font.size = Pt(20); r.font.color.rgb = RGBColor(0x1a, 0x34, 0x50)

    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = s.add_run("2026년 7월 17일 입국 기준 — 단계별 액션플랜")
    r2.font.size = Pt(12); r2.font.color.rgb = RGBColor(0xd4, 0xaf, 0x37)

    doc.add_paragraph()
    i = doc.add_paragraph(); i.alignment = WD_ALIGN_PARAGRAPH.CENTER
    i.add_run(f"작성일: 2026년 6월 11일\n대표이사: 이희성 (F-4 Overseas Korean Visa)\n투자자: U.S. Investor (D-8-1, 100% Equity Stake)").font.size = Pt(10)
    doc.add_page_break()

    add_heading_styled(doc, "개요", 1)
    add_body(doc,
        f"이 문서는 2026년 7월 17일 한국에 입국하는 F-4 비자 소지자 이희성 대표이사가 "
        f"8월 1일까지 법인설립을 완료하고, 이후 여행업 등록 및 모든 인허가를 완료하기까지의 "
        f"모든 단계를 시간순으로 정리한 실행 매뉴얼입니다.\n\n"
        f"핵심 전제:\n"
        f"• 투자자: 미국인, 자본금 {CAPITAL_INVESTOR}원 전액 출자 (지분 100%), D-8-1 비자 취득 예정\n"
        f"• 대표이사: 이희성, F-4 비자, 비주주 대표이사 (급여 수령)\n"
        f"• 법인형태: 외국인투자기업 주식회사 (FIPA 적용)\n"
        f"• 사업유형: 국내외여행업 (기획여행사 / Planned Travel Operator)\n"
        f"• 소재지: 부산광역시 해운대구\n"
        f"• 자본금: 총 {CAPITAL_TOTAL}원 (법정 최소 {CAPITAL_MIN}원 상회)\n"
        f"• 영업보증보험: {INSURANCE_TOTAL}원 (기본 {INSURANCE_BASE}원 + 기획여행 추가 {INSURANCE_PLANNED}원)\n"
        f"• 월고정비: ₩{MONTHLY_TOTAL} (보증보험 ₩{MONTHLY_INSURANCE_PREM} + 배상책임보험 ₩{MONTHLY_LIABILITY} 포함)\n"
        f"• 목표: 8월 중순까지 법인설립 완료 → 9월 중순 여행업 등록 완료 → 9월 말 영업 개시"
    )

    add_heading_styled(doc, "전체 타임라인 요약", 1)
    add_table(doc,
        ["단계", "예상 소요 기간", "누적", "목표 완료일"],
        [
            ["사전 준비 (미국)", "입국 전 완료", "D-14", "7월 3일까지"],
            ["STEP 1: 사무실 계약", "1~3일", "D+3", "7월 20일"],
            ["STEP 2: 정관 작성·공증", "2~3일", "D+6", "7월 23일"],
            ["STEP 3: 외국인투자신고 + 자본금 입금", "3~7일", "D+13", "7월 30일"],
            ["STEP 4: 법인설립등기", "5~7일", "D+20", "8월 6일"],
            ["STEP 5: 사업자등록", "3~5일", "D+25", "8월 11일"],
            ["STEP 6: 국내외여행업 등록", "14~30일", "D+55", "9월 10일"],
            ["STEP 7: 영업보증보험 가입", "5~7일", "D+62", "9월 17일"],
            ["STEP 8: D-8 비자 투자자 지원", "30~60일", "D+120", "10월~11월"],
            ["영업 개시", "—", "—", "9월 (등록 완료 즉시)"],
        ]
    )

    steps = [
        ("STEP 0: 입국 전 사전 준비 (현재 ~ 7월 16일)",
         f"소요 기간: 입국 전까지\n\n"
         f"1. 투자자와 주주간 계약서 체결\n"
         f"   - 지분 비율: 투자자 100% (CEO는 비주주 대표이사)\n"
         f"   - 자본금: 총 {CAPITAL_TOTAL}원 (투자자 {CAPITAL_INVESTOR}원 전액, CEO는 비주주)\n"
         f"   - 대표이사: 이희성 / 이사(또는 감사): 투자자\n"
         f"   - 배당 정책, 의사결정 권한, 지분양도 제한 등 명시\n\n"
         f"2. 투자자 서류 준비 (미국)\n"
         f"   - 여권 사본 (6개월 이상 유효기간)\n"
         f"   - FBI 범죄경력증명서 → 국무부 Apostille (D-8 비자용, 4~6주 소요)\n"
         f"   - 투자자금 원천 증명 (은행 statement, 세금신고서 등)\n\n"
         f"3. 한국 법무사/행정사 사전 섭외\n"
         f"   - 부산 지역 법인등기 전문 법무사\n"
         f"   - 예상 비용: 50~80만원 (법인설립 패키지)\n\n"
         f"4. 사무실 임차 사전 리서치\n"
         f"   - 해운대구 내 소형 오피스 (월 50~80만원)\n"
         f"   - 부동산 중개업소 컨택, 매물 리스트 확보\n\n"
         f"5. 외국인투자신고 사전 상담\n"
         f"   - KOTRA 외국인투자종합상담센터(1600-7119) 전화 상담\n"
         f"   - 외환은행 기업금융 지점 사전 방문 예약"),

        ("STEP 1: 사무실 계약 (7월 17~20일, 1~3일)",
         f"소요 기간: 1~3일\n\n"
         f"1. 입국 (7월 17일): 인천공항 → 부산 이동 (KTX 약 2시간40분)\n\n"
         f"2. 해운대구 부동산 방문 (7월 18일)\n"
         f"   - 사전 확보한 매물 리스트 중심으로 2~3곳 실사\n"
         f"   - 필수 요건: 실제 사무공간 (가상오피스 불가), 해운대구 관할, 확정일자 가능한 정식 임대차계약\n"
         f"   - 예상 보증금: 500~1,000만원 / 예상 월세: 50~80만원\n\n"
         f"3. 임대차 계약 체결 (7월 19~20일)\n\n"
         f"☑️ 확보: 임대차계약서 사본, 확정일자, 사업장 도면"),

        ("STEP 2: 정관 작성 및 공증 (7월 21~23일, 2~3일)",
         f"소요 기간: 2~3일\n\n"
         f"1. 정관 작성 (법무사 또는 startbiz.go.kr)\n"
         f"   - 상호: 리조네이트 클럽 주식회사 (Resonate Club Inc.)\n"
         f"   - 목적: 국내외여행업, 여행 알선, 관광 컨설팅 등\n"
         f"   - 본점: 부산광역시 해운대구\n"
         f"   - 자본금: {CAPITAL_TOTAL}원 (보통주 {SHARES}주, 1주당 ₩{PRICE_PER_SHARE})\n"
         f"   - 발기인: 이희성 + [투자자명]\n"
         f"   - 임원: 대표이사 1인, 이사(또는 감사) 1인\n\n"
         f"2. 공증: 부산지방법원 동부지원 인근 공증사무소 (비용 약 10~15만원)\n\n"
         f"☑️ 확보: 공증된 정관, 발기인 인감증명서"),

        ("STEP 3: 외국인투자신고 + 자본금 입금 (7월 24~30일, 3~7일)",
         f"소요 기간: 3~7일\n\n"
         f"⚠️ 중요: 반드시 외국인투자신고를 먼저 한 후 자본금을 입금해야 합니다. 순서가 바뀌면 과태료 및 세제혜택 상실 위험.\n\n"
         f"1. 외국인투자신고 (7월 24~25일)\n"
         f"   - 신고처: KOTRA 외국인투자종합상담센터 또는 외환은행 본점\n"
         f"   - 제출: 외국인투자신고서, 투자자 여권 사본, 주소 증명, 사업계획서, 정관 공증본\n"
         f"   - 처리 기간: 통상 1~2일\n\n"
         f"2. 자본금 계좌 개설 (7월 25~26일): 외환은행 또는 국민은행 '외국인투자 전용 계좌'\n\n"
         f"3. 자본금 입금 (7월 27~28일)\n"
         f"   - 투자자: {CAPITAL_INVESTOR}원 입금\n"
         f"   - CEO: 비주주 대표이사 — 자본금 납입 없음\n"
         f"   - 은행에서 '주식납입금 보관증명서' 발급\n\n"
         f"4. 외국인투자기업 등록 (7월 29~30일): KOTRA/은행에서 등록증 발급\n\n"
         f"☑️ 확보: 외국인투자기업등록증명서, 주식납입금 보관증명서"),

        ("STEP 4: 법인설립등기 (7월 31일 ~ 8월 6일, 5~7일)",
         f"소요 기간: 5~7일\n\n"
         f"1. 등기 신청서류 준비 (7월 31일 ~ 8월 1일)\n"
         f"   - 법무사에게 서류 일임 추천 (패키지 50~80만원)\n"
         f"   - 필요 서류: 정관 공증본, 발기인 인감증명서+주민등록등본, 주식납입금 보관증명서, 대표이사 취임승낙서, 이사 취임승낙서, 등기신청서+등록면허세\n\n"
         f"2. 등기소 접수 (8월 1~3일): 부산지방법원 동부지원 등기과 또는 iros.go.kr 전자등기\n\n"
         f"3. 등기완료 확인 (8월 6일 전후): 법인등기부등본 발급, 법인 인감카드, 법인 통장 개설\n\n"
         f"☑️ 확보: 법인등기부등본, 법인인감카드, 법인 통장"),

        ("STEP 5: 사업자등록 (8월 7~11일, 3~5일)",
         f"소요 기간: 3~5일\n\n"
         f"1. 사업자등록 신청 (8월 7~8일): 해운대세무서 방문 또는 hometax.go.kr 온라인\n"
         f"   - 필요 서류: 법인등기부등본, 임대차계약서 사본, 정관 사본, 법인인감증명서, 대표이사 신분증\n\n"
         f"2. 사업자등록증 수령 (8월 11일 전후)\n"
         f"   - 업종: 여행사업 (관광진흥법) / 관광여행사 / 일반과세자\n\n"
         f"☑️ 확보: 사업자등록증"),

        ("STEP 6: 국내외여행업 등록 (8월 12일 ~ 9월 10일, 14~30일)",
         f"소요 기간: 2~4주 (구청 및 한국관광공사 심사)\n\n"
         f"⚠️ 중요: 가장 긴 리드타임. 서류 누락 시 반려 → 재심사 2~4주 추가.\n\n"
         f"1. 구청 제출 서류 준비 (8월 12~14일)\n"
         f"   - 관광진흥법 시행규칙 제9조 서식: 여행업 등록신청서, 사업계획서, 법인등기부등본, 사업자등록증 사본, 임대차계약서 사본, 자본금 증명, 영업보증보험 가입 증명서, 국외여행 인솔자 자격증 사본, 수수료 5만원\n\n"
         f"2. 해운대구청 접수 (8월 14~18일): 관광과/문화관광과 방문\n\n"
         f"3. 심사 기간 (2~4주): 구청 1차 심사 → 필요 시 한국관광공사 기술 검토\n\n"
         f"4. 등록증 수령 (9월 10일 전후)\n\n"
         f"☑️ 확보: 국내외여행업 등록증"),

        ("STEP 7: 영업보증보험 가입 (8월 12일 ~ 9월 17일, 병행 진행)",
         f"소요 기간: SGI 심사 5~7일 (STEP 6과 병행 가능)\n\n"
         f"국내외여행업 + 기획여행(패키지) 기준 보증금액: 총 {INSURANCE_TOTAL}원\n"
         f"  → 기본 보증: {INSURANCE_BASE}원 (국내외여행업)\n"
         f"  → 기획여행 추가 보증: {INSURANCE_PLANNED}원\n"
         f"     (관광진흥법 시행규칙 별표3 비고 — 기본보증 외에 추가로 가입)\n\n"
         f"⚠️ 법령 근거: '국내외여행업 또는 종합여행업을 하는 여행업자 중에서 기획여행을 실시하려는 자는 "
         f"해당 업종에 따른 보증보험등에 가입하고 유지하는 것 외에 추가로 기획여행에 따른 보증보험등에 가입하여야 한다.'\n\n"
         f"1. 서울보증보험(SGI) 상담 예약 (8월 12일)\n"
         f"   - 부산지점: 부산광역시 연제구 법원로\n"
         f"   - 필요 서류: 법인등기부등본, 사업자등록증, 사업계획서, 재무제표/자본금 증명, 대표이사 신용정보 동의서\n\n"
         f"2. 보증보험 청약 (8월 14~15일)\n"
         f"   - 방식 A: 전액 지급보증 (예치금 0원, 연 보험료 0.5~1.5%)\n"
         f"     → 2.3억원 × 1% = 연 약 230만원 (월 약 {MONTHLY_INSURANCE_PREM}원)\n"
         f"   - 방식 B: 부분예치 + 지급보증 (심사 통과 어려울 시)\n"
         f"     → 예: 5,000만원 예치 + 1.8억원 지급보증\n\n"
         f"3. 보증서 발급 (8월 19~22일)\n\n"
         f"4. 배상책임보험 별도 가입 (월 약 {MONTHLY_LIABILITY}원, 관광진흥법상 의무)\n\n"
         f"☑️ 확보: 영업보증보험 증권 2건 (기본 + 기획여행 추가), 배상책임보험 증권"),

        ("STEP 8: 투자자 D-8 비자 지원 및 후속 조치 (9월 ~ )",
         f"소요 기간: 30~60일\n\n"
         f"9월부터 법인은 정상 영업 가능. 투자자 D-8 비자는 별도 진행.\n\n"
         f"1. 투자자에게 필요 서류 송부: 외국인투자기업등록증명서, 법인등기부등본(영문), 사업자등록증(영문), 주주명부, 정관, 사업계획서(영문), 사무실 임대차계약서\n\n"
         f"2. 투자자 → 주한미국대사관에서 D-8-1 비자 신청 (처리 기간: 5~10 영업일)\n\n"
         f"3. 투자자 입국 → 외국인등록증(ARC) 발급 (입국 후 90일 이내)\n\n"
         f"4. 추가 등록: 관광협회 가입, 국외여행 인솔자 등록, 여행자보험 상품 계약\n\n"
         f"※ D-8 비자 유효기간: {D8_VALIDITY}\n"
         f"※ 영주권(F-5) 경로: {F5_REQ}"),
    ]

    add_steps(doc, steps)

    add_heading_styled(doc, "종합 체크리스트", 1)
    add_table(doc,
        ["#", "완료 조건", "담당", "기한"],
        [
            ["1", "투자자 주주간 계약서 서명 완료", "CEO+투자자", "7월 10일"],
            ["2", "FBI 범죄증명서 Apostille 완료", "투자자", "7월 10일"],
            ["3", "사무실 임대차계약 + 확정일자", "CEO", "7월 20일"],
            ["4", "정관 공증 완료", "CEO", "7월 23일"],
            ["5", "외국인투자신고 접수 완료", "CEO(투자자 위임)", "7월 25일"],
            ["6", f"자본금 {CAPITAL_TOTAL}원 전액 입금", "투자자+CEO", "7월 28일"],
            ["7", "외국인투자기업 등록증 수령", "CEO", "7월 30일"],
            ["8", "법인설립등기 완료", "CEO+법무사", "8월 6일"],
            ["9", "사업자등록증 수령", "CEO", "8월 11일"],
            ["10", f"영업보증보험 증권 발급 ({INSURANCE_TOTAL}원)", "CEO", "8월 22일"],
            ["11", "배상책임보험 가입", "CEO", "8월 22일"],
            ["12", "국내외여행업 등록증 수령", "CEO", "9월 10일"],
            ["13", "D-8 비자 신청 서류 송부", "CEO", "9월 15일"],
            ["14", "영업 개시", "CEO", "9월 중순"],
        ]
    )

    add_body(doc, "\n---\n본 문서는 2026년 6월 기준 정보로 작성되었습니다. 실제 진행 시 법무사 및 세무사와 상담하십시오.")
    path = OUTDIR + "Resonate_Club_CEO_F4_Action_Plan_KO.docx"
    doc.save(path)
    print(f"[OK] Doc 2 KO: {path}")
    return path


# ══════════════════════════════════════════════════════════════
# DOC 2 — ENGLISH
# ══════════════════════════════════════════════════════════════
def create_doc2_en():
    doc = Document()
    set_narrow_margins(doc)

    doc.add_paragraph(); doc.add_paragraph()
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("RESONATE CLUB INC.\nCEO Action Plan — F-4 Visa Holder")
    r.bold = True; r.font.size = Pt(20); r.font.color.rgb = RGBColor(0x1a, 0x34, 0x50)

    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = s.add_run("Arrival July 17, 2026 — Step-by-Step Execution Timeline")
    r2.font.size = Pt(12); r2.font.color.rgb = RGBColor(0xd4, 0xaf, 0x37)

    doc.add_paragraph()
    i = doc.add_paragraph(); i.alignment = WD_ALIGN_PARAGRAPH.CENTER
    i.add_run(f"Date: June 11, 2026\nCEO: Hee Sung Lee (F-4 Overseas Korean Visa)\nInvestor: U.S. Investor (D-8-1, 100% Equity Stake)").font.size = Pt(10)
    doc.add_page_break()

    add_heading_styled(doc, "Overview", 1)
    add_body(doc,
        f"This document is a step-by-step execution manual for CEO Hee Sung Lee, arriving in Korea "
        f"on July 17, 2026 on an F-4 Overseas Korean Visa. It covers every action required to "
        f"establish the corporation by August 1 and complete all travel agency licensing by September.\n\n"
        f"Key Parameters:\n"
        f"• Investor: U.S. citizen/entity contributing KRW {CAPITAL_INVESTOR} (100% equity), D-8-1 visa\n"
        f"• CEO: Hee Sung Lee, F-4 visa, non-shareholder Representative Director (salaried)\n"
        f"• Entity: Foreign-Invested Joint-Stock Company under FIPA\n"
        f"• Business: Domestic & International Travel Agency (Planned Travel Operator)\n"
        f"• Location: Haeundae-gu, Busan, Republic of Korea\n"
        f"• Capital: KRW {CAPITAL_TOTAL} total (exceeds statutory minimum of KRW {CAPITAL_MIN})\n"
        f"• Guarantee Insurance: KRW {INSURANCE_TOTAL} (Base KRW {INSURANCE_BASE} + Planned Travel KRW {INSURANCE_PLANNED})\n"
        f"• Monthly Fixed Costs: KRW {MONTHLY_TOTAL} (incl. insurance KRW {MONTHLY_INSURANCE_PREM} + liability KRW {MONTHLY_LIABILITY})\n"
        f"• Goal: Corporation registered by Aug 6, travel license by Sep 10, operations begin late September"
    )

    add_heading_styled(doc, "Master Timeline Summary", 1)
    add_table(doc,
        ["Step", "Est. Duration", "Cumulative", "Target Date"],
        [
            ["Pre-Arrival Prep (US)", "Complete before arrival", "D-14", "By Jul 3"],
            ["STEP 1: Office Lease", "1–3 days", "D+3", "Jul 20"],
            ["STEP 2: Articles of Incorporation + Notarization", "2–3 days", "D+6", "Jul 23"],
            ["STEP 3: FDI Notification + Capital Deposit", "3–7 days", "D+13", "Jul 30"],
            ["STEP 4: Corporation Registration", "5–7 days", "D+20", "Aug 6"],
            ["STEP 5: Business Registration", "3–5 days", "D+25", "Aug 11"],
            ["STEP 6: Travel Agency License", "14–30 days", "D+55", "Sep 10"],
            ["STEP 7: Guarantee Insurance", "5–7 days", "D+62", "Sep 17"],
            ["STEP 8: Investor D-8 Visa Support", "30–60 days", "D+120", "Oct–Nov"],
            ["Operations Begin", "—", "—", "Late Sep (upon license)"],
        ]
    )

    steps_en = [
        ("STEP 0: Pre-Arrival Preparation (Now – Jul 16)",
         f"Duration: Before arrival\n\n"
         f"1. Sign Shareholder Agreement with investor\n"
         f"   - Equity: Investor 100% (CEO is non-shareholder Representative Director)\n"
         f"   - Capital: KRW {CAPITAL_TOTAL} (Investor KRW {CAPITAL_INVESTOR}, CEO is non-shareholder)\n"
         f"   - CEO: Hee Sung Lee / Director(or Auditor): Investor\n"
         f"   - Include: dividend policy, decision rights, transfer restrictions\n\n"
         f"2. Investor documents (US)\n"
         f"   - Passport copy (6+ months validity)\n"
         f"   - FBI background check → US Dept. of State Apostille (4–6 weeks; needed for D-8)\n"
         f"   - Source-of-funds documentation (bank statements, tax returns)\n\n"
         f"3. Engage Korean Certified Legal Administrator (Certified Legal Administrator)\n"
         f"   - Busan-based corporation registration specialist\n"
         f"   - Estimated cost: KRW 500,000–800,000 (incorporation package)\n\n"
         f"4. Office lease pre-research\n"
         f"   - Small office in Haeundae-gu (KRW 500K–800K/month)\n"
         f"   - Contact real estate agents, compile listing shortlist\n\n"
         f"5. FDI pre-consultation\n"
         f"   - Call KOTRA Invest KOREA Contact Center (1600-7119)\n"
         f"   - Schedule visit to foreign exchange bank corporate branch"),

        ("STEP 1: Office Lease (Jul 17–20, 1–3 days)",
         f"Duration: 1–3 days\n\n"
         f"1. Arrival (Jul 17): Incheon Airport → Busan via KTX (approx. 2h40m)\n\n"
         f"2. Property visits (Jul 18): 2–3 shortlisted offices in Haeundae-gu\n"
         f"   - Requirements: physical office (no virtual/coworking), Haeundae-gu jurisdiction, formal lease with fixed date stamp\n"
         f"   - Expected deposit: KRW 5M–10M / Monthly rent: KRW 500K–800K\n\n"
         f"3. Lease signing (Jul 19–20)\n\n"
         f"[CHECK] Lease agreement copy, Fixed Date Stamp, Office floor plan"),

        ("STEP 2: Articles of Incorporation + Notarization (Jul 21–23, 2–3 days)",
         f"Duration: 2–3 days\n\n"
         f"1. Draft Articles (Certified Legal Administrator or startbiz.go.kr)\n"
         f"   - Company: Resonate Club Inc. (Resonate Club Inc.)\n"
         f"   - Purpose: Domestic & International Travel Agency, travel consulting\n"
         f"   - Office: Haeundae-gu, Busan\n"
         f"   - Capital: KRW {CAPITAL_TOTAL} ({SHARES} common shares, KRW {PRICE_PER_SHARE}/share)\n"
         f"   - Founders: Hee Sung Lee + [Investor Name]\n"
         f"   - Officers: 1 Representative Director, 1 Director (or Auditor)\n\n"
         f"2. Notarization: Notary office near Busan District Court Eastern Branch (cost ~KRW 100K–150K)\n\n"
         f"[CHECK] Notarized Articles, Founder seal certificates"),

        ("STEP 3: FDI Notification + Capital Deposit (Jul 24–30, 3–7 days)",
         f"Duration: 3–7 days\n\n"
         f"⚠️ CRITICAL: FDI notification MUST precede capital transfer. Reversing the order triggers penalties and loss of tax benefits.\n\n"
         f"1. File Foreign Investment Notification (Jul 24–25)\n"
         f"   - At: KOTRA Invest KOREA or foreign exchange bank headquarters\n"
         f"   - Submit: FDI notification form, investor passport, proof of address, business plan, notarized Articles\n"
         f"   - Processing: typically 1–2 days\n\n"
         f"2. Open capital deposit account (Jul 25–26): Foreign-investment dedicated account at KEB Hana or KB Kookmin Bank\n\n"
         f"3. Transfer capital (Jul 27–28)\n"
         f"   - Investor: KRW {CAPITAL_INVESTOR}\n"
         f"   - CEO: Non-shareholder Representative Director — no capital contribution\n"
         f"   - Bank issues Share Subscription Deposit Certificate\n\n"
         f"4. FDI Enterprise Registration (Jul 29–30): KOTRA/bank issues FDI certificate\n\n"
         f"[CHECK] FDI Enterprise Registration Certificate, Share Subscription Deposit Certificate"),

        ("STEP 4: Corporation Registration (Jul 31 – Aug 6, 5–7 days)",
         f"Duration: 5–7 days\n\n"
         f"1. Prepare registration application (Jul 31 – Aug 1)\n"
         f"   - Recommend delegating to Certified Legal Administrator (package KRW 500K–800K)\n"
         f"   - Required: notarized Articles, founder seal certificates + resident registration, deposit certificate, CEO acceptance letter, director acceptance letter, application + registration tax\n\n"
         f"2. File at court (Aug 1–3): Busan District Court Eastern Branch registry or iros.go.kr electronic filing\n\n"
         f"3. Registration complete (by Aug 6): Corporate Registry Extract, corporate seal card, corporate bank account\n\n"
         f"[CHECK] Corporate Registry Extract, Corporate Seal Card, Corporate Bank Account"),

        ("STEP 5: Business Registration (Aug 7–11, 3–5 days)",
         f"Duration: 3–5 days\n\n"
         f"1. Submit application (Aug 7–8): Haeundae Tax Office or hometax.go.kr online\n"
         f"   - Required: Corporate Registry Extract, lease agreement, Articles, corporate seal certificate, CEO ID\n\n"
         f"2. Receive Business Registration Certificate (by Aug 11)\n"
         f"   - Business type: Travel Agency (Tourism Promotion Act) / General Taxpayer\n\n"
         f"[CHECK] Business Registration Certificate"),

        ("STEP 6: Travel Agency License (Aug 12 – Sep 10, 14–30 days)",
         f"Duration: 2–4 weeks (Haeundae-gu Office + Korea Tourism Organization review)\n\n"
         f"⚠️ Longest lead time. Document omissions = rejection + another 2–4 weeks.\n\n"
         f"1. Prepare submission package (Aug 12–14)\n"
         f"   - Tourism Promotion Act Enforcement Rules Art. 9: application form, business plan, Corporate Registry Extract, Business Registration Certificate, lease, capital proof, guarantee insurance certificate, tour escort qualification, fee KRW 50K\n\n"
         f"2. Submit to Haeundae-gu Office (Aug 14–18): Tourism/Culture & Tourism Division\n\n"
         f"3. Review period (2–4 weeks): 1st review by district office → technical review by KTO if needed\n\n"
         f"4. License issued (by Sep 10)\n\n"
         f"[CHECK] Domestic & International Travel Agency License"),

        ("STEP 7: Guarantee Insurance (Aug 12 – Sep 17, parallel with STEP 6)",
         f"Duration: SGI review 5–7 days (parallel with STEP 6)\n\n"
         f"Domestic & International Travel Agency + Planned Travel bond: KRW {INSURANCE_TOTAL} total\n"
         f"  → Base bond: KRW {INSURANCE_BASE}\n"
         f"  → Planned Travel surcharge: KRW {INSURANCE_PLANNED} (Tourism Act Enforcement Rules Schedule 3)\n\n"
         f"Legal basis: 'A travel agency operator conducting planned travel shall, in addition to maintaining "
         f"the guarantee insurance for their license type, additionally subscribe to guarantee insurance "
         f"for planned travel.'\n\n"
         f"1. SGI consultation (Aug 12): Busan branch, Yeonje-gu\n"
         f"   - Required: Corporate Registry Extract, Business Registration Certificate, business plan, financial statements, CEO credit consent\n\n"
         f"2. Insurance application (Aug 14–15)\n"
         f"   - Option A: Full surety bond (zero deposit, annual premium 0.5–1.5%)\n"
         f"     → KRW 230M × 1% = ~KRW 2.3M/year (~KRW {MONTHLY_INSURANCE_PREM}/month)\n"
         f"   - Option B: Partial deposit + surety (if full approval difficult)\n"
         f"     → Example: KRW 50M deposit + KRW 180M surety bond\n\n"
         f"3. Bond certificates issued (Aug 19–22)\n\n"
         f"4. Liability insurance: separate mandatory policy (~KRW {MONTHLY_LIABILITY}/month)\n\n"
         f"[CHECK] 2 Guarantee Insurance Certificates (base + planned travel), Liability Insurance Certificate"),

        ("STEP 8: Investor D-8 Visa Support + Follow-up (Sep onward)",
         f"Duration: 30–60 days\n\n"
         f"Corporation can operate from September. Investor D-8 visa is a separate parallel process.\n\n"
         f"1. Send required documents to investor: FDI certificate, Corp. Registry Extract (English), Business Registration (English), Shareholder Registry, Articles, Business Plan (English), Office lease\n\n"
         f"2. Investor applies for D-8-1 visa at Korean Embassy/Consulate in US (5–10 business days)\n\n"
         f"3. Investor enters Korea → Alien Registration Card (ARC) (within 90 days)\n\n"
         f"4. Additional registrations: Korea Tourism Association membership, Tour Escort registration, Traveler's Insurance product\n\n"
         f"D-8 Visa validity: {D8_VALIDITY}\n"
         f"Permanent Residency (F-5) path: {F5_REQ}"),
    ]

    add_steps(doc, steps_en)

    add_heading_styled(doc, "Comprehensive Checklist", 1)
    add_table(doc,
        ["#", "Condition", "Owner", "Deadline"],
        [
            ["1", "Shareholder agreement signed", "CEO+Investor", "Jul 10"],
            ["2", "FBI background check Apostille complete", "Investor", "Jul 10"],
            ["3", "Office lease signed + fixed date stamp", "CEO", "Jul 20"],
            ["4", "Articles notarized", "CEO", "Jul 23"],
            ["5", "FDI notification filed", "CEO (with investor POA)", "Jul 25"],
            ["6", f"Capital KRW {CAPITAL_TOTAL} fully deposited", "Investor+CEO", "Jul 28"],
            ["7", "FDI Enterprise Registration Certificate received", "CEO", "Jul 30"],
            ["8", "Corporation registration complete", "CEO+Certified Legal Administrator", "Aug 6"],
            ["9", "Business Registration Certificate received", "CEO", "Aug 11"],
            ["10", f"Guarantee Insurance bonds issued (KRW {INSURANCE_TOTAL})", "CEO", "Aug 22"],
            ["11", "Liability insurance policy active", "CEO", "Aug 22"],
            ["12", "Travel Agency License received", "CEO", "Sep 10"],
            ["13", "D-8 visa application package sent to investor", "CEO", "Sep 15"],
            ["14", "Operations commence", "CEO", "Late Sep"],
        ]
    )

    add_body(doc, "\n---\nThis document is based on information available as of June 2026. Consult a Korean Certified Legal Administrator and Certified Tax Accountant before proceeding.")
    path = OUTDIR + "Resonate_Club_CEO_F4_Action_Plan_EN.docx"
    doc.save(path)
    print(f"[OK] Doc 2 EN: {path}")
    return path


# ══════════════════════════════════════════════════════════════
# DOC 3 — KOREAN
# ══════════════════════════════════════════════════════════════
def create_doc3_ko():
    doc = Document()
    set_narrow_margins(doc)

    doc.add_paragraph(); doc.add_paragraph()
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("리조네이트 클럽 (Resonate Club)\n미국 현지 파트너사 사업계획 및 타임라인")
    r.bold = True; r.font.size = Pt(18); r.font.color.rgb = RGBColor(0x1a, 0x34, 0x50)

    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = s.add_run("한국→미국 프리미엄 골프 투어 · 현지 벤더 관리 · 공동 운영")
    r2.font.size = Pt(11); r2.font.color.rgb = RGBColor(0xd4, 0xaf, 0x37)

    doc.add_paragraph()
    i = doc.add_paragraph(); i.alignment = WD_ALIGN_PARAGRAPH.CENTER
    i.add_run(f"작성일: 2026년 6월 11일\n대표이사: 이희성 | 브랜드: 결 투어 (Resonate Tour)\n웹사이트: resonateclub.com / resonatetour.com").font.size = Pt(10)
    doc.add_page_break()

    add_heading_styled(doc, "1. 사업 개요", 1)
    add_body(doc,
        f"리조네이트 클럽 주식회사(Resonate Club Inc.)는 부산 해운대구에 본사를 둔 프리미엄 여행 기획사로, "
        f"'결 투어(Resonate Tour)' 브랜드로 한국인 대상 미국 프리미엄 골프 투어를 주력으로 운영합니다.\n\n"
        f"핵심 사업 모델:\n"
        f"• 한국인 고소득 골퍼 대상 → 캘리포니아 프리미엄 골프 투어\n"
        f"• 올인클루시브 패키지 (항공·숙박·골프·식사·교통·트립릴)\n"
        f"• 소규모 프라이빗 그룹 (2~8인) 맞춤형 운영\n"
        f"• 원가 +40% 마크업 투명 가격 정책 (고객에게는 원가 비공개)\n"
        f"• 추후 미국인→한국 인바운드 럭셔리 투어로 확장\n\n"
        f"회사 현황:\n"
        f"• 자본금: {CAPITAL_TOTAL}원 (법정 최소 {CAPITAL_MIN}원 상회, 투자자 100%)\n"
        f"• 영업보증보험: {INSURANCE_TOTAL}원 (기본 {INSURANCE_BASE}원 + 기획여행 추가 {INSURANCE_PLANNED}원)\n"
        f"• 월고정비: ₩{MONTHLY_TOTAL} (사무실 ₩{MONTHLY_OFFICE} + 보증보험 ₩{MONTHLY_INSURANCE_PREM} + 배상책임 ₩{MONTHLY_LIABILITY} + 회계 ₩{MONTHLY_ACCOUNTING} + 통신 ₩{MONTHLY_TELECOM} + 잡비 ₩{MONTHLY_MISC})\n\n"
        f"파트너사의 역할:\n"
        f"• 미국 현지에 상주하며 모든 현지 벤더(골프장, 호텔, 교통, 식당 등) 관리\n"
        f"• 신규 골프 코스 발굴 및 계약\n"
        f"• 투어 운영 중 현지 상황 대응 (실시간 문제 해결)\n"
        f"• 고객 만족도 관리 및 피드백 수집\n"
        f"• 현지 마케팅 협력 (인스타그램 콘텐츠 등)"
    )

    add_heading_styled(doc, "2. 법인 설립 타임라인 (파트너사 참고용)", 1)
    add_table(doc,
        ["시기", "이벤트", "의미"],
        [
            ["2026년 7월 17일", "CEO 한국 입국 (F-4 비자)", "법인 설립 시작"],
            ["2026년 7월 말", f"외국인투자신고 + 자본금 {CAPITAL_TOTAL}원 입금", f"자본금 확보 (법정 최소 {CAPITAL_MIN}원 상회)"],
            ["2026년 8월 초", "법인설립등기 완료", "Resonate Club Inc. 공식 출범"],
            ["2026년 8월 중순", "사업자등록 완료", "세금계산서 발행 가능"],
            ["2026년 9월 초", "국내외여행업 등록 완료", "여행 상품 판매 가능"],
            ["2026년 9월 중순", f"영업보증보험 가입 완료 ({INSURANCE_TOTAL}원)", "법적 요건 100% 충족, 영업 개시"],
            ["2026년 10월", "첫 투어 팀 출발 (목표)", "실전 운영 시작"],
        ]
    )

    add_heading_styled(doc, "3. 주요 여행 패키지 구성", 1)
    add_heading_styled(doc, "3.1 한국→미국 골프 투어 (주력)", 2)
    add_table(doc,
        ["패키지명", "기간", "라운드", "판매가 (USD/인)", "주요 코스"],
        [
            ["Trump Premier", "6박7일", "4라운드", "$9,752", "Trump National LA, Sandpiper GC, Hidden Valley GC, The Crossings"],
            ["Trump Elite", "5박6일", "3라운드", "$7,741", "Trump National LA, Sandpiper GC, Hidden Valley GC"],
        ]
    )

    add_heading_styled(doc, "3.2 향후 확장: 미국→한국 인바운드", 2)
    add_table(doc,
        ["패키지명", "기간", "판매가 (USD/인)", "주요 콘텐츠"],
        [
            ["Korea Golf Premier", "11박12일", "$27,300", "Nine Bridges, South Cape, Jack Nicklaus GC"],
            ["Korea Discovery", "12박13일", "$27,580", "K-Culture + Golf + Michelin Dining"],
        ]
    )

    add_heading_styled(doc, "3.3 VIP 이벤트 (기획 중)", 2)
    add_table(doc,
        ["이벤트명", "일정", "가격", "특징"],
        [["Moving Day VVIP Invitational", "2026년 6월 3-8일", "$14,400/인", "8인 한정, 호스트 Aimee Cho"]]
    )

    add_heading_styled(doc, "4. 파트너사 세부 역할 및 책임", 1)
    roles = [
        ("4.1 골프 코스 벤더 관리",
         "• 캘리포니아 내 프리미엄 골프 코스 발굴 및 계약\n"
         "  - 현재 협의 중: Trump National LA, Sandpiper GC, Hidden Valley GC, The Crossings, Alisal River Course\n"
         "  - 신규 발굴 목표: Pebble Beach, Torrey Pines, Pelican Hill\n"
         "• 티타임 확보: 시즌별·요일별 안정적 티타임 블록 계약\n"
         "• 그린피 협상: 단체 할인(4인+), 리플레이 할인, 시즌 특가\n"
         "• 코스 상태 모니터링: 에어레이션 일정, 잔디 상태, 공사 정보 사전 파악"),
        ("4.2 호텔 벤더 관리",
         "• 5성급 호텔 객실 블록 계약\n"
         "  - 주요 협력: Beverly Wilshire, Terranea Resort, Waldorf Astoria Beverly Hills\n"
         "  - 시즌별 단가 협상, 조식 포함 여부, 취소/변경 조건\n"
         "• 고객 등급별 호텔 옵션 다양화 (5성급 + 4성급 대안)\n"
         "• 체크인/체크아웃 원활한 진행 관리"),
        ("4.3 교통·물류 관리",
         "• 공항 픽업/샌딩 서비스: ICN↔LAX 연계\n"
         "• 투어 전용 차량: 12인승 프리미엄 밴 + 전담 드라이버\n"
         "• 코스 간 이동 동선 최적화 (로드매니저 역할)\n"
         "• 헬리콥터/요트 등 VIP 특별 이동수단 수배\n"
         "• 골프백 운송 관리 (항공 수하물 + 지상 이동)"),
        ("4.4 식음료·레스토랑 관리",
         "• 미쉐린 레스토랑 예약 및 코스 메뉴 조율\n"
         "• 한식 옵션이 있는 레스토랑 리스트 유지\n"
         "• 코스 내 식사(턴하우스/클럽하우스) 사전 주문\n"
         "• 고객 식이 제한(알러지 등) 대응\n"
         "• LA 한인타운 식당과의 제휴 (가이드+고객 식사)"),
        ("4.5 현지 가이드 관리",
         "• 한국어·영어 능통 가이드 섭외 및 관리\n"
         "• 가이드 교육: 코스 정보, 고객 응대 매뉴얼\n"
         "• 투어별 가이드 배정 및 일정 조율\n"
         "• 비상 상황 발생 시 의사소통 채널 유지"),
        ("4.6 마케팅·콘텐츠 협력",
         "• 골프 코스 촬영 (프로모션용 사진·영상)\n"
         "• 인스타그램/유튜브 콘텐츠용 현지 현장 영상 제공\n"
         "• 시네마틱 트립 릴 촬영 지원\n"
         "• 현지 골프 이벤트·토너먼트 정보 공유"),
    ]
    for title_text, body_text in roles:
        add_heading_styled(doc, title_text, 2)
        for line in body_text.split("\n"):
            if line.strip():
                add_body(doc, line.strip())

    add_heading_styled(doc, "5. 커뮤니케이션 및 업무 체계", 1)
    add_table(doc,
        ["항목", "내용"],
        [
            ["주요 연락 채널", "KakaoTalk (real-time), Email (formal), Google Drive (shared docs)"],
            ["정기 미팅", "Weekly video call (Zoom/Google Meet) — itinerary coordination, issue review"],
            ["긴급 연락", "24/7 KakaoTalk + shared Korea/US phone numbers"],
            ["문서 공유", "Google Drive: vendor contracts, pricing sheets, tour schedules, client info (encrypted)"],
            ["보고 체계", "Monthly: vendor status / Quarterly: financial reconciliation / Ad-hoc: new course/hotel proposals"],
            ["의사 결정", "Routine: Partner autonomy / >$3,000 spend: CEO approval / Strategic: joint decision"],
        ]
    )

    add_heading_styled(doc, "6. 재무 구조 및 수익 배분", 1)
    add_heading_styled(doc, "6.1 비용 구조 (투어 1팀 기준, 4인)", 2)
    add_table(doc,
        ["비용 항목", "추정 원가 (USD)", "비고"],
        [
            ["항공권 (ICN↔LAX 왕복)", "$6,000–8,000", "4인 합계, 이코노미 기준"],
            ["골프 그린피 (4라운드)", "$5,000–8,000", "코스 구성에 따라"],
            ["호텔 숙박 (6박)", "$6,000–12,000", "5성급, 2객실"],
            ["전용 밴 + 드라이버", "$3,000–4,500", "7일 전일정"],
            ["가이드 비용", "$2,000–3,000", "Korean-speaking guide"],
            ["식사 (전체)", "$2,500–4,000", "Restaurants + on-course dining"],
            ["기타 (트립릴 등)", "$500–1,000", "Filming, editing, souvenirs"],
            ["총 원가", "$25,000–40,500", ""],
            ["판매가 (4인)", "$35,000–56,700", "Cost + 40% markup"],
            ["마진 (4인)", "$10,000–16,200", "28–40% margin"],
        ]
    )

    add_heading_styled(doc, "6.2 수익 배분 모델 (제안)", 2)
    add_body(doc,
        "파트너사 수익 배분 구조 제안:\n\n"
        "1. 기본 운영 수수료: 투어당 순이익의 20~30%\n"
        "   - 골프장·호텔·교통·식당 벤더 관리 보상\n"
        "   - 실제 투어 진행 모니터링 및 문제 해결\n\n"
        "2. 신규 벤더 발굴 인센티브: 첫 계약 성사 시 원가 절감액의 30%\n\n"
        "3. 시즌 성과 보너스: 연간 목표 초과 달성 시 (연 12팀 초과 시 팀당 추가 보너스)\n\n"
        "4. 인바운드(미국→한국) 확장 시: 인바운드 전담 파트너십 재계약\n\n"
        "※ 구체적인 비율과 조건은 상호 협의하여 결정합니다."
    )

    add_heading_styled(doc, "7. 연간 운영 목표 (2026년 9월 ~ 2027년 8월)", 1)
    add_table(doc,
        ["구분", "Year 1 목표", "월 평균", "비고"],
        [
            ["투어 팀 수", "12~18팀", "1~1.5팀/월", "보수적 시나리오"],
            ["팀당 평균 인원", "4명", "—", "소규모 프라이빗"],
            ["연간 매출", "약 6억~9억원", "—", "$440K~$660K"],
            ["연간 순이익", "약 1.5억~2.5억원", "—", "마진율 25~35%"],
            ["신규 코스 계약", "3~5곳", "—", "캘리포니아 중심"],
            ["신규 호텔 계약", "2~3곳", "—", "5성급 위주"],
            ["SNS 팔로워", "5,000+", "—", "인스타그램 + 유튜브"],
            ["고객 만족도", "4.8/5.0", "—", "리뷰 평점 목표"],
        ]
    )

    add_heading_styled(doc, "8. 파트너사 즉시 준비 사항 (2026년 7월~8월)", 1)
    add_body(doc,
        "1. 사업자 등록 (미국): Sole Proprietorship 또는 LLC 설립\n\n"
        "2. 벤더 사전 접촉\n"
        "   - 기존 관계 활용하여 2026년 하반기 티타임·객실 사전 확보\n"
        "   - Trump National LA: 가을 시즌 티타임 블록 논의\n"
        "   - 호텔: 9~11월 객실 블록 계약 추진\n\n"
        "3. 현장 답사: 주요 코스 재방문 → 사진·영상 콘텐츠 확보, 호텔 객실·레스토랑 실사, 코스 간 이동 시간 실측\n\n"
        "4. 비상 매뉴얼 작성: 고객 부상/질병 시 병원 리스트, 항공편 지연/결항 대응 절차, 천재지변 대응 플랜\n\n"
        "5. 첫 투어 리허설: 9월 말~10월 초 첫 투어 전 코스·호텔·식당 최종 확인, CEO와 합동 리허설 진행"
    )

    add_heading_styled(doc, "9. 맺음말", 1)
    add_body(doc,
        "리조네이트 클럽은 단순한 여행사가 아닌, '한국과 미국을 잇는 프리미엄 경험 설계사'입니다.\n\n"
        "파트너사는 미국 현지에서 이 비전을 실현하는 핵심 축입니다. "
        "한국 본사가 상품 기획·마케팅·고객 관리를 담당하는 동안, "
        "파트너사는 최고의 현지 경험을 제공할 수 있도록 모든 인프라를 관리합니다.\n\n"
        "이 파트너십이 성공하기 위한 핵심 요소:\n"
        "• 투명한 커뮤니케이션 — 문제는 숨기지 않고 즉시 공유\n"
        "• 고객 중심 사고 — '고객에게 잊지 못할 경험을'이라는 공동 미션\n"
        "• 지속적 개선 — 매 투어 후 피드백 반영하여 업그레이드\n"
        "• 상호 신뢰 — 정직한 정산, 공정한 수익 배분\n\n"
        "함께 만들어갈 이 여정에 대한 기대와 확신을 담아, 본 계획서를 제안합니다.\n\n"
        "대표이사 이희성 드림\nresonatetour.com / resonateclub.com"
    )

    path = OUTDIR + "Resonate_Club_US_Partner_Brief_KO.docx"
    doc.save(path)
    print(f"[OK] Doc 3 KO: {path}")
    return path


# ══════════════════════════════════════════════════════════════
# DOC 3 — ENGLISH
# ══════════════════════════════════════════════════════════════
def create_doc3_en():
    doc = Document()
    set_narrow_margins(doc)

    doc.add_paragraph(); doc.add_paragraph()
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("RESONATE CLUB\nU.S. Partner Business Plan & Timeline")
    r.bold = True; r.font.size = Pt(18); r.font.color.rgb = RGBColor(0x1a, 0x34, 0x50)

    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = s.add_run("Korea→USA Premium Golf Tours · Local Vendor Management · Joint Operations")
    r2.font.size = Pt(11); r2.font.color.rgb = RGBColor(0xd4, 0xaf, 0x37)

    doc.add_paragraph()
    i = doc.add_paragraph(); i.alignment = WD_ALIGN_PARAGRAPH.CENTER
    i.add_run(f"Date: June 11, 2026\nCEO: Hee Sung Lee | Brand: Resonate Tour (Resonate Tour)\nWeb: resonateclub.com / resonatetour.com").font.size = Pt(10)
    doc.add_page_break()

    add_heading_styled(doc, "1. Business Overview", 1)
    add_body(doc,
        f"Resonate Club Inc. is a premium travel planning company headquartered in Haeundae-gu, Busan. "
        f"Under the 'Resonate Tour' brand, it specializes in premium U.S. golf tours for Korean travelers.\n\n"
        f"Core Business Model:\n"
        f"• Target: High-net-worth Korean golfers → California premium golf tours\n"
        f"• All-inclusive packages (flights, hotels, golf, dining, transport, cinematic trip reel)\n"
        f"• Small private groups (2–8 persons), fully customized\n"
        f"• Cost + 40% transparent markup (cost never disclosed to clients)\n"
        f"• Future expansion: USA→Korea inbound luxury tours\n\n"
        f"Company Snapshot:\n"
        f"• Capital: KRW {CAPITAL_TOTAL} (statutory minimum KRW {CAPITAL_MIN}; Investor 100%)\n"
        f"• Guarantee Insurance: KRW {INSURANCE_TOTAL} (Base KRW {INSURANCE_BASE} + Planned Travel KRW {INSURANCE_PLANNED})\n"
        f"• Monthly Fixed Costs: KRW {MONTHLY_TOTAL} (Office KRW {MONTHLY_OFFICE} + Guarantee Ins. KRW {MONTHLY_INSURANCE_PREM} + Liability KRW {MONTHLY_LIABILITY} + Accounting KRW {MONTHLY_ACCOUNTING} + Telecom KRW {MONTHLY_TELECOM} + Misc KRW {MONTHLY_MISC})\n\n"
        f"Partner Role:\n"
        f"• Based in the U.S., managing all local vendors (golf courses, hotels, transport, restaurants)\n"
        f"• Discovering and contracting new golf courses\n"
        f"• Real-time problem resolution during tours\n"
        f"• Customer satisfaction management and feedback collection\n"
        f"• Local marketing collaboration (Instagram content, etc.)"
    )

    add_heading_styled(doc, "2. Corporation Setup Timeline (Partner Reference)", 1)
    add_table(doc,
        ["Timing", "Event", "Significance"],
        [
            ["Jul 17, 2026", "CEO arrives Korea (F-4 visa)", "Setup begins"],
            ["Late Jul 2026", f"FDI notification + Capital deposit (KRW {CAPITAL_TOTAL})", "Capital secured"],
            ["Early Aug 2026", "Corporation registration complete", "Resonate Club Inc. officially established"],
            ["Mid Aug 2026", "Business registration complete", "Tax invoices can be issued"],
            ["Mid-Sep 2026", "Travel Agency License issued", "Tour products can be sold"],
            ["Mid Sep 2026", f"Guarantee insurance (KRW {INSURANCE_TOTAL}) active", "100% legally compliant; operations begin"],
            ["Oct 2026", "First tour group departure (target)", "Live operations commence"],
        ]
    )

    add_heading_styled(doc, "3. Core Tour Packages", 1)
    add_heading_styled(doc, "3.1 Korea→USA Golf Tours (Primary)", 2)
    add_table(doc,
        ["Package", "Duration", "Rounds", "Price (USD/person)", "Key Courses"],
        [
            ["Trump Premier", "6N/7D", "4 rounds", "$9,752", "Trump National LA, Sandpiper GC, Hidden Valley GC, The Crossings"],
            ["Trump Elite", "5N/6D", "3 rounds", "$7,741", "Trump National LA, Sandpiper GC, Hidden Valley GC"],
        ]
    )

    add_heading_styled(doc, "3.2 Future: USA→Korea Inbound", 2)
    add_table(doc,
        ["Package", "Duration", "Price (USD/person)", "Highlights"],
        [
            ["Korea Golf Premier", "11N/12D", "$27,300", "Nine Bridges, South Cape, Jack Nicklaus GC"],
            ["Korea Discovery", "12N/13D", "$27,580", "K-Culture + Golf + Michelin Dining"],
        ]
    )

    add_heading_styled(doc, "3.3 VIP Events (Planning Stage)", 2)
    add_table(doc,
        ["Event", "Dates", "Price", "Feature"],
        [["Moving Day VVIP Invitational", "Jun 3-8, 2026", "$14,400/person", "8 guests only, Host Aimee Cho"]]
    )

    add_heading_styled(doc, "4. Partner Roles & Responsibilities", 1)
    roles_en = [
        ("4.1 Golf Course Vendor Management",
         "• Discover and contract premium California golf courses\n"
         "  - Current/negotiating: Trump National LA, Sandpiper GC, Hidden Valley GC, The Crossings, Alisal River Course\n"
         "  - Target acquisitions: Pebble Beach, Torrey Pines, Pelican Hill\n"
         "• Secure stable seasonal/weekly tee time blocks\n"
         "• Negotiate green fees: group discounts (4+), replay discounts, seasonal specials\n"
         "• Monitor course conditions: aeration schedules, turf status, construction alerts"),
        ("4.2 Hotel Vendor Management",
         "• Block-book 5-star hotel rooms\n"
         "  - Key partners: Beverly Wilshire, Terranea Resort, Waldorf Astoria Beverly Hills\n"
         "  - Seasonal rate negotiation, breakfast inclusion, cancellation/change terms\n"
         "• Diversify hotel tiers (5-star + 4-star alternatives)\n"
         "• Seamless check-in/check-out management"),
        ("4.3 Transportation & Logistics",
         "• Airport pickup/drop-off: ICN↔LAX coordination\n"
         "• Tour vehicle: 12-passenger premium van + dedicated driver\n"
         "• Route optimization between courses (road manager function)\n"
         "• VIP transport: helicopter, yacht arrangements\n"
         "• Golf bag logistics (airline baggage + ground transport)"),
        ("4.4 Dining & Restaurant Management",
         "• Michelin restaurant reservations and menu coordination\n"
         "• Maintain list of Korean-cuisine-friendly restaurants\n"
         "• Pre-order on-course meals (turn house/clubhouse)\n"
         "• Accommodate dietary restrictions (allergies, etc.)\n"
         "• Partnerships with LA Koreatown restaurants (guide + guest meals)"),
        ("4.5 Local Guide Management",
         "• Recruit and manage Korean-English bilingual guides\n"
         "• Guide training: course information, client service protocols\n"
         "• Tour-specific guide assignments and schedule coordination\n"
         "• Emergency communication channels"),
        ("4.6 Marketing & Content Collaboration",
         "• Golf course photography/videography (promotional content)\n"
         "• Provide on-location footage for Instagram/YouTube\n"
         "• Support cinematic trip reel production\n"
         "• Share local golf events and tournament information"),
    ]
    for title_text, body_text in roles_en:
        add_heading_styled(doc, title_text, 2)
        for line in body_text.split("\n"):
            if line.strip():
                add_body(doc, line.strip())

    add_heading_styled(doc, "5. Communication & Operations Framework", 1)
    add_table(doc,
        ["Item", "Details"],
        [
            ["Primary Channels", "KakaoTalk (real-time), Email (formal), Google Drive (shared documents)"],
            ["Regular Meetings", "Weekly video call (Zoom/Google Meet) — itinerary coordination, issue review"],
            ["Emergency Contact", "24/7 KakaoTalk + shared Korea/US phone numbers"],
            ["Document Sharing", "Google Drive: vendor contracts, pricing sheets, tour schedules, client info (encrypted)"],
            ["Reporting", "Monthly: vendor status / Quarterly: financial reconciliation / Ad-hoc: new proposals"],
            ["Decision Making", "Routine: Partner autonomy / >$3,000: CEO approval / Strategic: joint decision"],
        ]
    )

    add_heading_styled(doc, "6. Financial Structure & Revenue Sharing", 1)
    add_heading_styled(doc, "6.1 Cost Structure (Per Tour, 4 Guests)", 2)
    add_table(doc,
        ["Cost Item", "Est. Cost (USD)", "Note"],
        [
            ["Flights (ICN↔LAX round-trip)", "$6,000–8,000", "4 persons, economy"],
            ["Green Fees (4 rounds)", "$5,000–8,000", "Course-dependent"],
            ["Hotel (6 nights)", "$6,000–12,000", "5-star, 2 rooms"],
            ["Private Van + Driver", "$3,000–4,500", "7 days, full schedule"],
            ["Guide Fee", "$2,000–3,000", "Korean-speaking guide"],
            ["Meals (all)", "$2,500–4,000", "Restaurants + on-course"],
            ["Other (Film Reel, etc.)", "$500–1,000", "Filming, editing, souvenirs"],
            ["Total Cost", "$25,000–40,500", ""],
            ["Sell Price (4 guests)", "$35,000–56,700", "Cost + 40% markup"],
            ["Margin (4 guests)", "$10,000–16,200", "28–40% margin"],
        ]
    )

    add_heading_styled(doc, "6.2 Revenue Share Model (Proposed)", 2)
    add_body(doc,
        "Proposed revenue share structure:\n\n"
        "1. Base Operating Commission: 20–30% of net profit per tour\n"
        "   - Compensation for golf, hotel, transport, restaurant vendor management\n"
        "   - Real-time tour monitoring and issue resolution\n\n"
        "2. New Vendor Discovery Incentive: 30% of first-year cost savings from new contracts\n\n"
        "3. Seasonal Performance Bonus: For exceeding annual targets (per-team bonus above 12 teams/year)\n\n"
        "4. Inbound (USA→Korea) Expansion: Separate inbound partnership agreement upon launch\n\n"
        "※ Specific ratios and conditions to be mutually agreed."
    )

    add_heading_styled(doc, "7. Annual Operating Targets (Sep 2026 – Aug 2027)", 1)
    add_table(doc,
        ["Metric", "Year 1 Target", "Monthly Avg", "Note"],
        [
            ["Tour Groups", "12–18 groups", "1–1.5/month", "Conservative scenario"],
            ["Avg Group Size", "4 persons", "—", "Small private groups"],
            ["Annual Revenue", "~KRW 600M–900M", "—", "~$440K–$660K"],
            ["Annual Net Profit", "~KRW 150M–250M", "—", "25–35% margin"],
            ["New Course Contracts", "3–5 courses", "—", "California focus"],
            ["New Hotel Contracts", "2–3 hotels", "—", "5-star priority"],
            ["Social Media Followers", "5,000+", "—", "Instagram + YouTube"],
            ["Customer Satisfaction", "4.8/5.0", "—", "Review rating target"],
        ]
    )

    add_heading_styled(doc, "8. Immediate Partner Action Items (Jul–Aug 2026)", 1)
    add_body(doc,
        "1. U.S. Business Registration: Sole Proprietorship or LLC\n\n"
        "2. Vendor Pre-Contact\n"
        "   - Leverage existing relationships to pre-secure H2 2026 tee times and room blocks\n"
        "   - Trump National LA: discuss fall season tee time blocks\n"
        "   - Hotels: pursue Sep–Nov room block contracts\n\n"
        "3. Site Inspections: Revisit key courses for photo/video content, inspect hotel rooms/restaurants, measure inter-course drive times\n\n"
        "4. Emergency Manual: Hospital list for guest injury/illness, flight delay/cancellation protocols, natural disaster response plan\n\n"
        "5. First Tour Rehearsal: Final course/hotel/restaurant verification before late Sep/early Oct launch, joint rehearsal with CEO"
    )

    add_heading_styled(doc, "9. Closing", 1)
    add_body(doc,
        "Resonate Club is not just a travel agency — it is a premium experience curator bridging Korea and the United States.\n\n"
        "The Partner is the critical on-the-ground axis that brings this vision to life. While the Korea headquarters "
        "handles product planning, marketing, and client management, the Partner manages all local infrastructure "
        "to deliver world-class experiences.\n\n"
        "Key success factors:\n"
        "• Transparent communication — surface problems immediately, never hide issues\n"
        "• Client-first mindset — shared mission of 'unforgettable experiences'\n"
        "• Continuous improvement — feedback-driven upgrades after every tour\n"
        "• Mutual trust — honest accounting, fair revenue sharing\n\n"
        "We present this plan with excitement and confidence in the journey ahead.\n\n"
        "Hee Sung Lee, CEO\nresonatetour.com / resonateclub.com"
    )

    path = OUTDIR + "Resonate_Club_US_Partner_Brief_EN.docx"
    doc.save(path)
    print(f"[OK] Doc 3 EN: {path}")
    return path


# ══════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════
if __name__ == "__main__":
    create_doc2_ko()
    create_doc2_en()
    create_doc3_ko()
    create_doc3_en()
    print(f"\n[DONE] 4 files generated in {OUTDIR}")
